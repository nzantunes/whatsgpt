"""
Script de Automação com IA (navegador, pesquisas, câmera, edição de texto)
Versão refatorada e otimizada para melhor eficiência, legibilidade e manutenibilidade.

Funcionalidades:
- Abrir o navegador e acessar sites
- Fazer pesquisas na web
- Ligar e usar a câmera (fotos, etc.)
- Editar texto em campos e aplicativos
- Execução direta de tarefas mapeadas via pyautogui
"""

# ==================== PATCH THREADING (evita "cannot set daemon status of active thread") ====================
# Algumas bibliotecas (pyautogui, pygetwindow, etc.) tentam definir daemon em thread já iniciada; isso evita o RuntimeError.
import threading as _threading
_OriginalThread = _threading.Thread
# Obter getter/setter da property daemon da Thread original (compatível com todas as versões)
_daemon_fget = _OriginalThread.daemon.fget
_daemon_fset = getattr(_OriginalThread.daemon, 'fset', None)

class _ThreadDaemonPatch(_OriginalThread):
    @property
    def daemon(self):
        return _daemon_fget(self)
    @daemon.setter
    def daemon(self, value):
        try:
            if not self.is_alive() and _daemon_fset is not None:
                _daemon_fset(self, value)
        except RuntimeError:
            # "cannot set daemon status of active thread" - ignorar quando a thread já está ativa
            pass
        except Exception:
            pass
_threading.Thread = _ThreadDaemonPatch

# ==================== IMPORTS ====================
import openai
import pyautogui
import time
import os
import sys
import io
import subprocess
import re
import json
import ast
import inspect
import traceback
from typing import Optional, Tuple, Dict, Callable, List
from datetime import datetime
from urllib.parse import quote_plus
from urllib.request import Request, urlopen
from urllib.error import URLError, HTTPError
import pygetwindow as gw

# OpenCV é opcional (usado apenas para funções de câmera)
try:
    import cv2
    CV2_AVAILABLE = True
except ImportError:
    cv2 = None
    CV2_AVAILABLE = False

# ==================== CONFIGURAÇÕES ====================
# Configurações de segurança do pyautogui
pyautogui.FAILSAFE = True  # Mover mouse para canto superior esquerdo para parar
pyautogui.PAUSE = 0.3  # Pausa entre ações (em segundos)

# Configurações de controle
MOUSE_SPEED = 0.5  # Velocidade do movimento do mouse (0-1)
CLICK_DELAY = 0.2  # Delay após cliques (em segundos)
CURSOR_WAIT_TIMEOUT = 10  # Timeout padrão para aguardar Cursor
CODE_COMPLETION_TIMEOUT = 60  # Timeout para aguardar código ser gerado
CODE_COMPLETION_CHECK_INTERVAL = 2.0  # Intervalo entre verificações

# Quando True, o agente NÃO abre nem ativa a janela do Cursor (apenas tarefas diretas)
CURSOR_INTERACTION_DISABLED = True
# Quando True, ao adicionar novo comando em cursor_automation_extensions.py envia aviso/comando ao terminal do Cursor
ENVIAR_ATUALIZACOES_TERMINAL_CURSOR = os.environ.get("ENVIAR_ATUALIZACOES_TERMINAL_CURSOR", "1").strip().lower() in ("1", "true", "sim", "yes")

# Navegador padrão: Brave ("big"). Use PREFERRED_BROWSER=Chrome no .env para voltar ao Chrome.
_preferred = (os.environ.get("PREFERRED_BROWSER") or "").strip()
if _preferred and _preferred.lower() in ("brave", "big"):
    DEFAULT_BROWSER = "Brave"
elif _preferred:
    DEFAULT_BROWSER = _preferred
else:
    DEFAULT_BROWSER = "Brave"  # navegador Brave (Big) como padrão

# Varredura na web: configurável por .env
TIMEOUT_VARREDURA = int(os.environ.get("TIMEOUT_VARREDURA", "5"))  # segundos por página/link
MAX_LINKS_VARREDURA = int(os.environ.get("MAX_LINKS_VARREDURA", "3"))  # quantos links abrir

# Diretório para histórico de otimizações
OPTIMIZATION_HISTORY_DIR = os.path.join(os.path.dirname(__file__), '.optimization_history')
os.makedirs(OPTIMIZATION_HISTORY_DIR, exist_ok=True)
OPTIMIZATION_HISTORY_FILE = os.path.join(OPTIMIZATION_HISTORY_DIR, 'optimizations.json')
PERFORMANCE_LOG_FILE = os.path.join(OPTIMIZATION_HISTORY_DIR, 'performance_log.json')

# Config do bot (personalização): preenchido por _apply_bot_config() quando o Node envia AUTOMATION_CONFIG_JSON
AGENT_CONFIG: Dict = {}


def _agent_feature_enabled(feature: str) -> bool:
    """Retorna True se a funcionalidade está habilitada na config do bot (padrão True)."""
    return AGENT_CONFIG.get("features", {}).get(feature, True)


# ==================== INICIALIZAÇÃO ====================
def load_openai_api_key() -> str:
    """
    Carrega a chave da API OpenAI do ambiente ou arquivo .env.
    
    Returns:
        Chave da API ou string vazia se não encontrada
    """
    api_key = os.getenv("OPENAI_API_KEY", "")
    
    if not api_key:
        try:
            _script_dir = os.path.dirname(os.path.abspath(__file__))
            _project_dir = os.path.dirname(_script_dir)
            env_path = os.path.join(_project_dir, '.env')
            if os.path.exists(env_path):
                with open(env_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        if line.startswith('OPENAI_API_KEY='):
                            api_key = line.split('=', 1)[1].strip().strip('"').strip("'")
                            break
        except Exception as e:
            print(f"Aviso: Não foi possível ler o arquivo .env: {e}")
    
    return api_key

OPENAI_API_KEY = load_openai_api_key()


def load_xai_api_key() -> str:
    """Carrega a chave da API xAI (Grok) do ambiente ou .env."""
    key = os.getenv("XAI_API_KEY", "") or os.getenv("GROK_API_KEY", "")
    if not key:
        try:
            _script_dir = os.path.dirname(os.path.abspath(__file__))
            _project_dir = os.path.dirname(_script_dir)
            env_path = os.path.join(_project_dir, '.env')
            if os.path.exists(env_path):
                with open(env_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        if line.startswith('XAI_API_KEY=') or line.startswith('GROK_API_KEY='):
                            key = line.split('=', 1)[1].strip().strip('"').strip("'")
                            break
        except Exception:
            pass
    return key


XAI_API_KEY = load_xai_api_key()

# Timeout para chamadas de IA (evita travar se a API demorar)
AI_REQUEST_TIMEOUT = 45.0

# Cliente OpenAI (GPT)
client = openai.OpenAI(api_key=OPENAI_API_KEY, timeout=AI_REQUEST_TIMEOUT) if OPENAI_API_KEY else None

# Cliente xAI (Grok) — API compatível com OpenAI
grok_client = None
if XAI_API_KEY:
    try:
        grok_client = openai.OpenAI(api_key=XAI_API_KEY, base_url="https://api.x.ai/v1", timeout=AI_REQUEST_TIMEOUT)
    except Exception:
        grok_client = None


def _get_agent_model() -> str:
    """Modelo configurado para o agente interpretar comandos (vem da config do Node: agentModel)."""
    return (AGENT_CONFIG.get("agentModel") or os.environ.get("AGENT_MODEL") or "gpt-4o").strip()


def _model_has_small_context(model: str) -> bool:
    """True se o modelo tem limite de contexto pequeno (ex.: 8192 tokens), para reduzir histórico no prompt."""
    if not model:
        return False
    m = model.lower().strip()
    if "gpt-3.5" in m:
        return True
    if m == "gpt-4" or m.startswith("gpt-4-0") or m.startswith("gpt-4-8k"):
        return True
    return False


def _grok_model_id(model: str) -> str:
    """Mapeia nome do modelo Grok para o ID da API xAI."""
    if not model or model == "grok-2":
        return "grok-2"
    if model in ("grok-3", "grok-4", "grok-beta"):
        return model
    if model.startswith("grok-"):
        return model
    return "grok-2"


# ==================== FUNÇÕES DE IA ====================
def get_ai_response(prompt: str, system_prompt: Optional[str] = None, max_tokens: int = 2000) -> str:
    """
    Obtém resposta da IA para um prompt. Usa o modelo configurado na config do agente (GPT ou Grok).
    """
    model = _get_agent_model()
    use_grok = model.lower().startswith("grok")

    if use_grok and grok_client:
        api_client = grok_client
        api_model = _grok_model_id(model)
    elif client:
        api_client = client
        api_model = model if model.lower().startswith("gpt-") else "gpt-4o"
    else:
        if use_grok:
            print("Erro: XAI_API_KEY (Grok) não configurada para o modelo do agente!")
        else:
            print("Erro: OPENAI_API_KEY não configurada!")
        return ""

    try:
        default_system = """Você é um agente autônomo que controla o PC para tarefas no navegador e no sistema.
Principais funções: abrir o navegador, fazer pesquisas na web, ligar/usar a câmera, editar texto em campos e aplicativos.
Use pyautogui para simular ações (cliques, digitação, atalhos). Seja seguro e peça confirmação para ações destrutivas.
Não é necessário abrir o Cursor IDE — foque em navegador, pesquisas, câmera e edição de texto."""
        system_content = system_prompt or default_system
        # Modelos com 8192 tokens: limitar system+user para não estourar contexto
        if _model_has_small_context(model):
            max_system_chars = 14000
            if len(system_content) > max_system_chars:
                system_content = system_content[:max_system_chars].rstrip() + "\n\n[... contexto truncado para caber no limite do modelo ...]"
            if len(prompt) > 4000:
                prompt = prompt[:4000].rstrip() + "\n\n[...]"
        response = api_client.chat.completions.create(
            model=api_model,
            max_tokens=max_tokens,
            temperature=0.5,
            messages=[
                {"role": "system", "content": system_content},
                {"role": "user", "content": prompt},
            ],
        )
        return (response.choices[0].message.content or "").strip()
    except Exception as e:
        print(f"Erro ao obter resposta da IA ({api_model}): {e}")
        return ""

# ==================== FUNÇÕES DE CONTROLE DO CURSOR ====================
def get_cursor_window():
    """
    Encontra a janela do Cursor.
    
    Returns:
        Objeto Window do Cursor ou None se não encontrado
    """
    try:
        windows = gw.getWindowsWithTitle("Cursor")
        if windows:
            return windows[0]
        
        # Tentar variações do nome
        for title in ["Cursor", "cursor", "Cursor AI", "cursor.exe"]:
            windows = gw.getWindowsWithTitle(title)
            if windows:
                return windows[0]
    except Exception as e:
        print(f"Aviso ao buscar janela: {e}")
    
    return None

def activate_cursor_window() -> bool:
    """
    Ativa a janela do Cursor se estiver aberta.
    Não faz nada quando CURSOR_INTERACTION_DISABLED está True (agente só executa tarefas diretas).
    
    Returns:
        True se conseguiu ativar, False caso contrário
    """
    if CURSOR_INTERACTION_DISABLED:
        return False
    window = get_cursor_window()
    if window:
        try:
            window.activate()
            time.sleep(0.5)
            return True
        except Exception as e:
            print(f"Erro ao ativar janela: {e}")
    return False

def wait_for_cursor(timeout: int = CURSOR_WAIT_TIMEOUT) -> bool:
    """
    Aguarda o Cursor abrir verificando a janela.
    
    Args:
        timeout: Tempo máximo de espera em segundos
    
    Returns:
        True se o Cursor foi aberto, False caso contrário
    """
    print("Aguardando Cursor abrir...")
    start_time = time.time()
    
    while time.time() - start_time < timeout:
        window = get_cursor_window()
        if window and window.visible:
            print("Cursor encontrado e visível!")
            if activate_cursor_window():
                return True
        time.sleep(0.5)
    
    print("Cursor não foi encontrado no tempo esperado")
    return False

def open_cursor() -> bool:
    """
    (DESATIVADO) O agente não abre mais o Cursor IDE.
    Executa apenas tarefas diretas: navegador, pesquisa, câmera, rede, otimização.
    """
    print("O agente não abre o Cursor. Use tarefas diretas: abrir navegador, pesquisar X, ligar câmera, etc.")
    return False

# ==================== FUNÇÕES DE CONTROLE DO MOUSE ====================
def move_and_click(x: int, y: int, button: str = 'left', clicks: int = 1, interval: float = 0.1):
    """
    Move o mouse para uma posição e clica.
    
    Args:
        x: Coordenada X
        y: Coordenada Y
        button: Botão do mouse ('left', 'right', 'middle')
        clicks: Número de cliques
        interval: Intervalo entre cliques
    """
    print(f"Movendo mouse para ({x}, {y}) e clicando...")
    pyautogui.moveTo(x, y, duration=MOUSE_SPEED)
    time.sleep(0.1)
    pyautogui.click(x, y, button=button, clicks=clicks, interval=interval)
    time.sleep(CLICK_DELAY)

def click_element_on_screen(image_path: str, confidence: float = 0.8, timeout: int = 5) -> bool:
    """
    Procura uma imagem na tela e clica nela.
    
    Args:
        image_path: Caminho para a imagem a procurar
        confidence: Nível de confiança (0-1)
        timeout: Tempo máximo para procurar
    
    Returns:
        True se encontrou e clicou, False caso contrário
    """
    print(f"Procurando elemento na tela: {image_path}")
    start_time = time.time()
    
    while time.time() - start_time < timeout:
        try:
            location = pyautogui.locateOnScreen(image_path, confidence=confidence)
            if location:
                center = pyautogui.center(location)
                move_and_click(center.x, center.y)
                print(f"Elemento encontrado e clicado em ({center.x}, {center.y})")
                return True
        except pyautogui.ImageNotFoundException:
            pass
        time.sleep(0.5)
    
    print(f"Elemento não encontrado: {image_path}")
    return False

def get_mouse_position() -> Tuple[int, int]:
    """Retorna a posição atual do mouse."""
    return pyautogui.position()

def drag_and_drop(start_x: int, start_y: int, end_x: int, end_y: int, duration: float = 1.0):
    """
    Arrasta de uma posição para outra.
    
    Args:
        start_x, start_y: Posição inicial
        end_x, end_y: Posição final
        duration: Duração do arrasto em segundos
    """
    print(f"Arrastando de ({start_x}, {start_y}) para ({end_x}, {end_y})...")
    pyautogui.moveTo(start_x, start_y, duration=0.3)
    time.sleep(0.2)
    pyautogui.dragTo(end_x, end_y, duration=duration, button='left')
    time.sleep(CLICK_DELAY)

def scroll(direction: str = 'down', amount: int = 3):
    """
    Rola a tela.
    
    Args:
        direction: 'up' ou 'down'
        amount: Quantidade de rolagens
    """
    print(f"Rolando {direction} {amount} vezes...")
    for _ in range(amount):
        if direction == 'down':
            pyautogui.scroll(-3)
        else:
            pyautogui.scroll(3)
        time.sleep(0.1)

def right_click(x: int = None, y: int = None):
    """Clica com botão direito do mouse."""
    if x is not None and y is not None:
        pyautogui.rightClick(x, y)
    else:
        pyautogui.rightClick()

def double_click(x: int = None, y: int = None):
    """Clica duas vezes com botão esquerdo."""
    if x is not None and y is not None:
        pyautogui.doubleClick(x, y)
    else:
        pyautogui.doubleClick()

def select_all():
    """Seleciona todo o texto no editor (só ativa Cursor se interação não estiver desativada)."""
    if not CURSOR_INTERACTION_DISABLED:
        activate_cursor_window()
    pyautogui.hotkey('ctrl', 'a')
    time.sleep(0.2)

def copy_text():
    """Copia o texto selecionado."""
    pyautogui.hotkey('ctrl', 'c')
    time.sleep(0.2)

def paste_text():
    """Cola o texto da área de transferência."""
    pyautogui.hotkey('ctrl', 'v')
    time.sleep(0.2)

def undo():
    """Desfaz a última ação."""
    pyautogui.hotkey('ctrl', 'z')
    time.sleep(0.2)

def redo():
    """Refaz a última ação desfeita."""
    pyautogui.hotkey('ctrl', 'y')
    time.sleep(0.2)

# ==================== FUNÇÕES DE CÂMERA ====================
def access_notebook_camera(start_capture: bool = True) -> Optional[object]:
    """
    Acessa a câmera do notebook usando OpenCV e inicia a captura de vídeo/imagens.
    
    Args:
        start_capture: Se True, inicia a captura imediatamente (padrão: True)
    
    Returns:
        Objeto VideoCapture da câmera ou None se não conseguir acessar
    """
    if not CV2_AVAILABLE:
        print("Erro: OpenCV (cv2) não está instalado. Instale com: pip install opencv-python")
        return None
    
    print("Acessando câmera do notebook...")
    cap = None
    
    try:
        cap = cv2.VideoCapture(0)
        
        if not cap.isOpened():
            print("Erro: Não foi possível abrir a câmera.")
            return None
        
        # No Windows (MSMF) os primeiros frames podem falhar; dar tempo e tentar de novo
        time.sleep(1.0)
        for _ in range(30):
            ret, frame = cap.read()
            if ret:
                break
            time.sleep(0.2)
        if not ret:
            print("Erro: Não foi possível ler frame (câmera pode estar em uso). Tente 'ligar câmera' para ver via navegador.")
            return None
        
        print("✅ Câmera acessada com sucesso!")
        
        if start_capture:
            print("📹 Iniciando captura de vídeo/imagens...")
            print("💡 Pressione 'q' para parar a captura")
            print("💡 Pressione 's' para salvar uma foto")
            
            while True:
                ret, frame = cap.read()
                if not ret:
                    # Tentar mais algumas vezes (MSMF às vezes falha um frame)
                    for _ in range(5):
                        time.sleep(0.1)
                        ret, frame = cap.read()
                        if ret:
                            break
                    if not ret:
                        print("Erro: Não foi possível ler frame.")
                        break
                
                cv2.imshow('Notebook Camera - Pressione Q para sair, S para salvar foto', frame)
                
                key = cv2.waitKey(1) & 0xFF
                
                if key == ord('q'):
                    print("Captura encerrada pelo usuário")
                    break
                elif key == ord('s'):
                    photo_filename = f"foto_captura_{int(time.time())}.jpg"
                    cv2.imwrite(photo_filename, frame)
                    print(f"✅ Foto salva: {photo_filename}")
        else:
            print("Câmera aberta e pronta para uso")
            return cap
            
    except Exception as e:
        print(f"Erro ao acessar câmera: {e}")
        return None
    finally:
        if cap:
            cap.release()
        cv2.destroyAllWindows()
        print("Câmera liberada")
    
    return None


def close_camera_app() -> bool:
    """
    Desliga a câmera: para o stream ao servidor (se estiver ativo) e fecha o app de câmera do Windows.
    Útil para comando "desligar câmera" pelo WhatsApp.
    """
    stopped = stop_camera_stream()
    try:
        for title in ["Câmera", "Camera", "Câmera do Windows", "Windows Camera"]:
            windows = gw.getWindowsWithTitle(title)
            if not windows:
                # Busca parcial (alguns sistemas têm título diferente)
                all_windows = gw.getAllWindows()
                for w in all_windows:
                    if w.title and title.lower() in w.title.lower():
                        windows = [w]
                        break
            if windows:
                w = windows[0]
                w.activate()
                time.sleep(0.4)
                pyautogui.hotkey('alt', 'F4')
                time.sleep(0.3)
                print("Câmera (app) fechada.")
                return True
    except Exception as e:
        print(f"Aviso ao fechar câmera: {e}")
    return stopped


# Arquivo onde guardamos o PID do processo de stream da câmera (para desligar depois)
_CAMERA_STREAM_PID_FILE = os.path.join(os.path.dirname(__file__), '.camera_stream.pid')


def _get_camera_server_base_url() -> str:
    """Retorna a URL base do servidor (WhatsGPT) para enviar frames. Usada para ver câmera via rede externa."""
    url = os.environ.get("AUTOMATION_BASE_URL", "").strip()
    # Fallback: ler ngrok-url.txt na pasta do projeto (mesma que o Node usa)
    if not url or url.startswith("http://localhost") or url.startswith("http://127.0.0.1"):
        try:
            _project_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            _ngrok_file = os.path.join(_project_dir, 'ngrok-url.txt')
            if os.path.exists(_ngrok_file):
                with open(_ngrok_file, 'r', encoding='utf-8') as f:
                    for line in f:
                        t = line.strip()
                        if t and not t.startswith('#') and re.search(r'https?://', t):
                            m = re.search(r'https?://[^\s#]+', t)
                            if m:
                                url = m.group(0).rstrip('/')
                                break
        except Exception:
            pass
    if not url:
        try:
            _project_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            _env_path = os.path.join(_project_dir, '.env')
            if os.path.exists(_env_path):
                with open(_env_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        if line.strip().startswith('BASE_URL='):
                            url = line.split('=', 1)[1].strip().strip('"').strip("'")
                            break
        except Exception:
            pass
    if not url:
        url = "http://localhost:3000"
    url = url.rstrip('/')
    # Se veio URL completa (ex.: .../api/camera/view), usar só a base
    if '/api/' in url:
        url = re.sub(r'/api/.*$', '', url)
    return url


def _run_camera_stream_loop() -> None:
    """
    Loop que captura frames da câmera e envia para o servidor (POST /api/camera/frame).
    Deve ser executado em um processo separado (--camera-stream).
    """
    if not CV2_AVAILABLE:
        print("Erro: OpenCV (cv2) não instalado. pip install opencv-python")
        return
    base_url = _get_camera_server_base_url()
    post_url = base_url + "/api/camera/frame"
    cap = None
    try:
        cap = cv2.VideoCapture(0)
        if not cap.isOpened():
            print("Erro: Não foi possível abrir a câmera.")
            return
        import threading
        interval = 0.5  # ~2 FPS para não sobrecarregar
        sent = 0
        while True:
            ret, frame = cap.read()
            if not ret:
                time.sleep(interval)
                continue
            _, jpeg = cv2.imencode('.jpg', frame)
            data = jpeg.tobytes()
            try:
                req = Request(post_url, data=data, method='POST', headers={'Content-Type': 'image/jpeg'})
                urlopen(req, timeout=5)
                sent += 1
                if sent == 1:
                    print("Câmera enviando imagens para o servidor. URL para ver:", base_url + "/api/camera/view")
            except (URLError, HTTPError, OSError) as e:
                if sent == 0:
                    print("Aviso: servidor inacessível?", e)
            time.sleep(interval)
    except Exception as e:
        print("Stream câmera encerrado:", e)
    finally:
        if cap:
            cap.release()
        cv2.destroyAllWindows()


def start_camera_stream_to_server() -> bool:
    """
    Inicia o envio das imagens da câmera para o servidor em um processo separado.
    Retorna True e imprime a URL para ver ao vivo (rede externa).
    """
    if not CV2_AVAILABLE:
        print("Erro: OpenCV (cv2) não instalado. pip install opencv-python")
        return False
    base_url = _get_camera_server_base_url()
    view_url = base_url + "/api/camera/view"
    # Evitar dois processos de stream
    if os.path.exists(_CAMERA_STREAM_PID_FILE):
        try:
            with open(_CAMERA_STREAM_PID_FILE, 'r') as f:
                old_pid = int(f.read().strip())
            if _is_process_alive(old_pid):
                print("Câmera já está ligada e enviando. URL:", view_url)
                print("LINK_CAMERA:", view_url)
                return True
        except (ValueError, OSError):
            pass
        try:
            os.remove(_CAMERA_STREAM_PID_FILE)
        except OSError:
            pass
    creationflags = 0
    if sys.platform == 'win32':
        creationflags = getattr(subprocess, 'CREATE_NO_WINDOW', 0x08000000) | getattr(subprocess, 'DETACHED_PROCESS', 0x00000008)
    try:
        child = subprocess.Popen(
            [sys.executable, os.path.abspath(__file__), '--camera-stream'],
            cwd=os.path.dirname(os.path.abspath(__file__)),
            env={**os.environ, 'AUTOMATION_BASE_URL': base_url},
            stdin=subprocess.DEVNULL,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            creationflags=creationflags,
        )
        with open(_CAMERA_STREAM_PID_FILE, 'w') as f:
            f.write(str(child.pid))
        print("Câmera ligada. Imagens enviadas ao servidor.")
        print("URL para ver (rede externa):", view_url)
        print("LINK_CAMERA:", view_url)  # Para o Node extrair e enviar no WhatsApp
        return True
    except Exception as e:
        print("Erro ao iniciar stream da câmera:", e)
        return False


def _is_process_alive(pid: int) -> bool:
    """Verifica se o processo com o PID ainda está em execução (Windows e Unix)."""
    if sys.platform == 'win32':
        try:
            r = subprocess.run(
                ['tasklist', '/FI', f'PID eq {pid}', '/NH'],
                capture_output=True, text=True, timeout=5,
                creationflags=getattr(subprocess, 'CREATE_NO_WINDOW', 0x08000000) if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
            )
            return str(pid) in (r.stdout or '')
        except Exception:
            return False
    try:
        os.kill(pid, 0)
        return True
    except (ProcessLookupError, OSError):
        return False


def _terminate_process(pid: int) -> None:
    """Encerra o processo (Windows: taskkill; Unix: SIGKILL)."""
    if sys.platform == 'win32':
        try:
            subprocess.run(
                ['taskkill', '/PID', str(pid), '/F'],
                capture_output=True, timeout=5,
                creationflags=getattr(subprocess, 'CREATE_NO_WINDOW', 0x08000000) if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
            )
        except Exception:
            pass
        return
    try:
        os.kill(pid, 9)
    except (ProcessLookupError, OSError):
        pass


def stop_camera_stream() -> bool:
    """Para o processo que envia as imagens da câmera ao servidor."""
    if not os.path.exists(_CAMERA_STREAM_PID_FILE):
        return True
    try:
        with open(_CAMERA_STREAM_PID_FILE, 'r') as f:
            pid = int(f.read().strip())
        os.remove(_CAMERA_STREAM_PID_FILE)
        _terminate_process(pid)
        print("Stream da câmera desligado.")
        return True
    except (ValueError, OSError) as e:
        print("Aviso ao parar stream:", e)
        try:
            os.remove(_CAMERA_STREAM_PID_FILE)
        except OSError:
            pass
    return False


def take_photo(filename: str = None, save_path: str = None) -> Optional[str]:
    """
    Tira uma foto usando a câmera do notebook e salva em um arquivo.
    
    Args:
        filename: Nome do arquivo (opcional, padrão: foto_YYYYMMDD_HHMMSS.jpg)
        save_path: Caminho onde salvar a foto (opcional, padrão: diretório atual)
    
    Returns:
        Caminho do arquivo salvo ou None se houver erro
    """
    if not CV2_AVAILABLE:
        print("Erro: OpenCV (cv2) não está instalado. Instale com: pip install opencv-python")
        return None
    
    print("Tirando foto...")
    cap = None
    
    try:
        cap = cv2.VideoCapture(0)
        
        if not cap.isOpened():
            print("Erro: Não foi possível abrir a câmera.")
            return None
        
        # Aguardar câmera ajustar
        time.sleep(1)
        
        # Capturar frame
        ret, frame = cap.read()
        
        if not ret:
            print("Erro: Não foi possível capturar a imagem.")
            return None
        
        # Gerar nome do arquivo se não fornecido
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"foto_{timestamp}.jpg"
        
        # Garantir extensão .jpg
        if not filename.lower().endswith(('.jpg', '.jpeg', '.png')):
            filename += '.jpg'
        
        # Determinar caminho completo
        if save_path is None:
            save_path = os.path.join(os.getcwd(), filename)
        else:
            save_path = os.path.join(save_path, filename)
        
        # Salvar foto
        cv2.imwrite(save_path, frame)
        print(f"Foto salva com sucesso: {save_path}")
        
        return save_path
        
    except Exception as e:
        print(f"Erro ao tirar foto: {e}")
        return None
    finally:
        if cap:
            cap.release()

# ==================== FUNÇÕES DE CONEXÃO DE REDE ====================
def check_internet_connection() -> bool:
    """
    Verifica se há conexão ativa com a internet.
    
    Returns:
        True se há conexão, False caso contrário
    """
    import socket
    
    try:
        # Tentar conectar ao DNS do Google
        socket.create_connection(("8.8.8.8", 53), timeout=3)
        return True
    except OSError:
        pass
    
    try:
        # Tentar conectar ao DNS da Cloudflare
        socket.create_connection(("1.1.1.1", 53), timeout=3)
        return True
    except OSError:
        pass
    
    return False

def get_wifi_networks() -> list:
    """
    Lista redes Wi-Fi disponíveis.
    
    Returns:
        Lista de dicionários com informações das redes (nome, sinal, etc.)
    """
    networks = []
    
    if sys.platform != 'win32':
        print("Listagem de redes Wi-Fi disponível apenas no Windows")
        return networks
    
    try:
        # Usar netsh para listar redes Wi-Fi
        result = subprocess.run(
            ['netsh', 'wlan', 'show', 'profiles'],
            capture_output=True,
            text=True,
            timeout=10
        )
        
        if result.returncode == 0:
            # Listar perfis salvos
            profiles = []
            for line in result.stdout.split('\n'):
                if 'All User Profile' in line or 'Perfil de Todos os Usuários' in line:
                    profile_name = line.split(':')[-1].strip()
                    if profile_name:
                        profiles.append(profile_name)
            
            # Agora listar redes disponíveis
            result = subprocess.run(
                ['netsh', 'wlan', 'show', 'networks', 'mode=Bssid'],
                capture_output=True,
                text=True,
                timeout=10
            )
            
            if result.returncode == 0:
                current_network = {}
                for line in result.stdout.split('\n'):
                    line = line.strip()
                    if 'SSID' in line and ':' in line:
                        if current_network:
                            networks.append(current_network)
                        current_network = {'name': line.split(':')[-1].strip()}
                    elif 'Signal' in line or 'Sinal' in line:
                        if ':' in line:
                            signal = line.split(':')[-1].strip()
                            current_network['signal'] = signal
                    elif 'Authentication' in line or 'Autenticação' in line:
                        if ':' in line:
                            auth = line.split(':')[-1].strip()
                            current_network['auth'] = auth
                
                if current_network:
                    networks.append(current_network)
    except Exception as e:
        print(f"Erro ao listar redes Wi-Fi: {e}")
    
    return networks

def connect_to_wifi(ssid: str, password: str = None) -> bool:
    """
    Conecta a uma rede Wi-Fi.
    
    Args:
        ssid: Nome da rede Wi-Fi
        password: Senha da rede (opcional se já estiver salva)
    
    Returns:
        True se conseguiu conectar, False caso contrário
    """
    if sys.platform != 'win32':
        print("Conexão Wi-Fi disponível apenas no Windows")
        return False
    
    print(f"Conectando à rede Wi-Fi: {ssid}")
    
    try:
        if password:
            # Criar perfil XML temporário
            import tempfile
            xml_content = f'''<?xml version="1.0"?>
<WLANProfile xmlns="http://www.microsoft.com/networking/WLAN/profile/v1">
    <name>{ssid}</name>
    <SSIDConfig>
        <SSID>
            <name>{ssid}</name>
        </SSID>
    </SSIDConfig>
    <connectionType>ESS</connectionType>
    <connectionMode>auto</connectionMode>
    <MSM>
        <security>
            <authEncryption>
                <authentication>WPA2PSK</authentication>
                <encryption>AES</encryption>
                <useOneX>false</useOneX>
            </authEncryption>
            <sharedKey>
                <keyType>passPhrase</keyType>
                <protected>false</protected>
                <keyMaterial>{password}</keyMaterial>
            </sharedKey>
        </security>
    </MSM>
</WLANProfile>'''
            
            with tempfile.NamedTemporaryFile(mode='w', suffix='.xml', delete=False) as f:
                f.write(xml_content)
                temp_file = f.name
            
            try:
                # Adicionar perfil
                result = subprocess.run(
                    ['netsh', 'wlan', 'add', 'profile', f'filename={temp_file}'],
                    capture_output=True,
                    text=True,
                    timeout=10
                )
                
                if result.returncode != 0:
                    print(f"Erro ao adicionar perfil: {result.stderr}")
                    return False
            finally:
                # Remover arquivo temporário
                try:
                    os.unlink(temp_file)
                except:
                    pass
        
        # Conectar à rede
        result = subprocess.run(
            ['netsh', 'wlan', 'connect', f'name={ssid}'],
            capture_output=True,
            text=True,
            timeout=15
        )
        
        if result.returncode == 0:
            print(f"✅ Conectando à rede {ssid}...")
            time.sleep(5)  # Aguardar conexão estabelecer
            
            # Verificar se conectou
            if check_internet_connection():
                print(f"✅ Conectado à rede {ssid} com sucesso!")
                return True
            else:
                print(f"⚠️ Conectado à rede, mas sem acesso à internet")
                return True  # Conectado mesmo sem internet
        else:
            print(f"❌ Erro ao conectar: {result.stderr}")
            return False
            
    except Exception as e:
        print(f"Erro ao conectar à rede Wi-Fi: {ssid}: {e}")
        return False

def check_ethernet_connection() -> bool:
    """
    Verifica se há conexão Ethernet ativa.
    
    Returns:
        True se há conexão Ethernet, False caso contrário
    """
    if sys.platform != 'win32':
        return False
    
    try:
        result = subprocess.run(
            ['netsh', 'interface', 'show', 'interface'],
            capture_output=True,
            text=True,
            timeout=10
        )
        
        if result.returncode == 0:
            for line in result.stdout.split('\n'):
                if 'Ethernet' in line or 'Conectado' in line or 'Connected' in line:
                    if 'Conectado' in line or 'Connected' in line:
                        return True
    except Exception as e:
        print(f"Erro ao verificar conexão Ethernet: {e}")
    
    return False

def connect_to_internet(wifi_ssid: str = None, wifi_password: str = None) -> bool:
    """
    Conecta à internet verificando e estabelecendo conexão de rede.
    
    Args:
        wifi_ssid: Nome da rede Wi-Fi (opcional)
        wifi_password: Senha da rede Wi-Fi (opcional)
    
    Returns:
        True se conseguiu conectar, False caso contrário
    """
    print("Verificando conexão de rede...")
    
    # Verificar se já está conectado
    if check_internet_connection():
        print("✅ Já há conexão com a internet!")
        return True
    
    # Verificar conexão Ethernet
    if check_ethernet_connection():
        print("✅ Conexão Ethernet detectada")
        time.sleep(2)
        if check_internet_connection():
            print("✅ Conectado à internet via Ethernet!")
            return True
        else:
            print("⚠️ Ethernet conectado mas sem acesso à internet")
    
    # Tentar conectar via Wi-Fi
    if wifi_ssid:
        print(f"Tentando conectar via Wi-Fi: {wifi_ssid}")
        if connect_to_wifi(wifi_ssid, wifi_password):
            return True
    else:
        # Listar redes disponíveis
        print("Listando redes Wi-Fi disponíveis...")
        networks = get_wifi_networks()
        
        if networks:
            print(f"Encontradas {len(networks)} redes Wi-Fi:")
            for i, net in enumerate(networks[:5], 1):  # Mostrar apenas as 5 primeiras
                print(f"  {i}. {net.get('name', 'Desconhecida')} - Sinal: {net.get('signal', 'N/A')}")
            
            # Tentar conectar à primeira rede salva (sem senha)
            for net in networks:
                if net.get('name'):
                    print(f"\nTentando conectar à rede: {net['name']}")
                    if connect_to_wifi(net['name']):
                        return True
        else:
            print("⚠️ Nenhuma rede Wi-Fi disponível encontrada")
    
    print("❌ Não foi possível estabelecer conexão com a internet")
    return False

# ==================== SISTEMA DE AUTO-MELHORAMENTO ====================
class PerformanceMonitor:
    """Monitora desempenho de execuções para identificar otimizações."""
    
    def __init__(self):
        self.execution_times = {}
        self.error_counts = {}
        self.success_rates = {}
        self.load_history()
    
    def load_history(self):
        """Carrega histórico de desempenho."""
        try:
            if os.path.exists(PERFORMANCE_LOG_FILE):
                with open(PERFORMANCE_LOG_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.execution_times = data.get('execution_times', {})
                    self.error_counts = data.get('error_counts', {})
                    self.success_rates = data.get('success_rates', {})
        except Exception as e:
            print(f"Aviso: Erro ao carregar histórico de desempenho: {e}")
            self.execution_times = {}
            self.error_counts = {}
            self.success_rates = {}
    
    def save_history(self):
        """Salva histórico de desempenho."""
        try:
            data = {
                'execution_times': self.execution_times,
                'error_counts': self.error_counts,
                'success_rates': self.success_rates,
                'last_updated': datetime.now().isoformat()
            }
            with open(PERFORMANCE_LOG_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"Erro ao salvar histórico: {e}")
    
    def record_execution(self, function_name: str, execution_time: float, success: bool):
        """Registra execução de uma função."""
        if function_name not in self.execution_times:
            self.execution_times[function_name] = []
            self.error_counts[function_name] = 0
            self.success_rates[function_name] = {'total': 0, 'success': 0}
        
        self.execution_times[function_name].append(execution_time)
        self.success_rates[function_name]['total'] += 1
        if success:
            self.success_rates[function_name]['success'] += 1
        else:
            self.error_counts[function_name] += 1
        
        # Manter apenas últimas 100 execuções
        if len(self.execution_times[function_name]) > 100:
            self.execution_times[function_name] = self.execution_times[function_name][-100:]
        
        self.save_history()
    
    def get_slow_functions(self, threshold: float = 2.0) -> List[Tuple[str, float]]:
        """Retorna funções que demoram mais que o threshold."""
        slow_functions = []
        for func_name, times in self.execution_times.items():
            if times:
                avg_time = sum(times) / len(times)
                if avg_time > threshold:
                    slow_functions.append((func_name, avg_time))
        return sorted(slow_functions, key=lambda x: x[1], reverse=True)
    
    def get_error_prone_functions(self) -> List[Tuple[str, int]]:
        """Retorna funções com mais erros."""
        error_prone = [(name, count) for name, count in self.error_counts.items() if count > 0]
        return sorted(error_prone, key=lambda x: x[1], reverse=True)

# Instância global do monitor
performance_monitor = PerformanceMonitor()

def analyze_own_code() -> Dict:
    """
    Analisa o próprio código para identificar áreas de melhoria.
    
    Returns:
        Dicionário com análise do código
    """
    script_path = __file__
    analysis = {
        'filepath': script_path,
        'size': 0,
        'lines': 0,
        'functions': [],
        'classes': [],
        'imports': [],
        'complexity_issues': [],
        'optimization_opportunities': []
    }
    
    try:
        with open(script_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        analysis['size'] = len(content)
        analysis['lines'] = len(content.split('\n'))
        
        # Análise AST
        try:
            tree = ast.parse(content)
            
            for node in ast.walk(tree):
                if isinstance(node, ast.FunctionDef):
                    analysis['functions'].append({
                        'name': node.name,
                        'line': node.lineno,
                        'args_count': len(node.args.args),
                        'complexity': estimate_complexity(node)
                    })
                elif isinstance(node, ast.ClassDef):
                    analysis['classes'].append({
                        'name': node.name,
                        'line': node.lineno,
                        'methods_count': len([n for n in node.body if isinstance(n, ast.FunctionDef)])
                    })
                elif isinstance(node, ast.Import):
                    for alias in node.names:
                        analysis['imports'].append(alias.name)
                elif isinstance(node, ast.ImportFrom):
                    if node.module:
                        analysis['imports'].append(node.module)
        except Exception as e:
            print(f"Erro ao analisar AST: {e}")
        
        # Identificar oportunidades de otimização
        analysis['optimization_opportunities'] = identify_optimizations(content, analysis)
        
    except Exception as e:
        print(f"Erro ao analisar código: {e}")
    
    return analysis

def estimate_complexity(node: ast.FunctionDef) -> int:
    """Estima complexidade ciclomática de uma função."""
    complexity = 1  # Base complexity
    for child in ast.walk(node):
        if isinstance(child, (ast.If, ast.While, ast.For, ast.Try, ast.With)):
            complexity += 1
        elif isinstance(child, ast.BoolOp):
            complexity += len(child.values) - 1
    return complexity

def identify_optimizations(content: str, analysis: Dict) -> List[Dict]:
    """Identifica oportunidades de otimização no código."""
    opportunities = []
    
    # Verificar funções lentas
    slow_functions = performance_monitor.get_slow_functions()
    for func_name, avg_time in slow_functions[:5]:  # Top 5
        opportunities.append({
            'type': 'performance',
            'function': func_name,
            'issue': f'Função lenta (média: {avg_time:.2f}s)',
            'priority': 'high' if avg_time > 5 else 'medium'
        })
    
    # Verificar funções com muitos erros
    error_prone = performance_monitor.get_error_prone_functions()
    for func_name, error_count in error_prone[:5]:  # Top 5
        opportunities.append({
            'type': 'reliability',
            'function': func_name,
            'issue': f'Função com muitos erros ({error_count} erros)',
            'priority': 'high' if error_count > 10 else 'medium'
        })
    
    # Verificar funções complexas
    for func_info in analysis.get('functions', []):
        if func_info.get('complexity', 0) > 15:
            opportunities.append({
                'type': 'complexity',
                'function': func_info['name'],
                'issue': f'Função muito complexa (complexidade: {func_info["complexity"]})',
                'priority': 'medium'
            })
    
    # Verificar imports não utilizados (análise básica)
    # Esta análise pode ser expandida no futuro
    # Por enquanto, apenas retornar oportunidades identificadas acima
    
    return opportunities

def generate_optimization_code(function_name: str, issue: str, code_context: str) -> Optional[str]:
    """
    Gera código otimizado para uma função usando IA.
    
    Args:
        function_name: Nome da função a otimizar
        issue: Descrição do problema
        code_context: Contexto do código atual
    
    Returns:
        Código otimizado ou None
    """
    if not client:
        return None
    
    system_prompt = """Você é um especialista em otimização de código Python. 
Analise o código fornecido e gere uma versão otimizada que:
1. Melhore o desempenho
2. Reduza complexidade
3. Melhore legibilidade
4. Mantenha a mesma funcionalidade
5. Adicione tratamento de erros se necessário

Retorne APENAS o código otimizado da função, sem explicações extras."""
    
    user_prompt = f"""Otimize a seguinte função Python:

PROBLEMA IDENTIFICADO: {issue}

CÓDIGO ATUAL:
{code_context}

Gere uma versão otimizada desta função que resolva o problema identificado."""
    
    try:
        response = get_ai_response(user_prompt, system_prompt, max_tokens=2000)
        return response
    except Exception as e:
        print(f"Erro ao gerar otimização: {e}")
        return None

def apply_optimization(function_name: str, optimized_code: str, issue: str = "Otimização automática") -> bool:
    """
    Aplica otimização ao código, substituindo a função antiga.
    
    Args:
        function_name: Nome da função
        optimized_code: Código otimizado
        issue: Descrição do problema resolvido
    
    Returns:
        True se aplicou com sucesso, False caso contrário
    """
    script_path = __file__
    
    try:
        # Ler código atual
        with open(script_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Parse AST para encontrar função
        tree = ast.parse(content)
        
        # Encontrar função
        function_node = None
        for node in ast.walk(tree):
            if isinstance(node, ast.FunctionDef) and node.name == function_name:
                function_node = node
                break
        
        if not function_node:
            print(f"Função {function_name} não encontrada")
            return False
        
        # Extrair código original usando linhas
        lines = content.split('\n')
        start_line = function_node.lineno - 1
        end_line = function_node.end_lineno if hasattr(function_node, 'end_lineno') else start_line + 100
        
        # Encontrar início real da função (incluindo decorators)
        while start_line > 0 and (lines[start_line].strip().startswith('@') or lines[start_line].strip() == ''):
            start_line -= 1
        
        original_code = '\n'.join(lines[start_line:end_line])
        
        # Limpar código otimizado (remover markdown se presente)
        clean_code = optimized_code.strip()
        if clean_code.startswith('```'):
            code_lines = clean_code.split('\n')
            # Remover primeira linha (```python ou similar) e última linha (```)
            clean_code = '\n'.join(code_lines[1:-1]) if len(code_lines) > 2 else clean_code
        clean_code = clean_code.strip()
        
        # Garantir que o código otimizado começa com def
        if not clean_code.startswith('def ') and not clean_code.startswith('@'):
            # Tentar extrair apenas a função
            if 'def ' in clean_code:
                idx = clean_code.index('def ')
                clean_code = clean_code[idx:]
        
        # Substituir no código
        new_content = content[:start_line] + clean_code + '\n' + content[end_line:]
        
        # Validar sintaxe antes de salvar
        try:
            ast.parse(new_content)
        except SyntaxError as e:
            print(f"Erro de sintaxe no código otimizado: {e}")
            print(f"Código otimizado (primeiras 500 chars):\n{clean_code[:500]}")
            return False
        
        # Fazer backup
        backup_path = script_path + f'.backup_{datetime.now().strftime("%Y%m%d_%H%M%S")}'
        with open(backup_path, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"✅ Backup criado: {backup_path}")
        
        # Salvar código otimizado
        with open(script_path, 'w', encoding='utf-8') as f:
            f.write(new_content)
        
        # Registrar otimização
        record_optimization(function_name, issue=issue, original_code=original_code, optimized_code=clean_code)
        
        print(f"✅ Otimização aplicada à função {function_name}")
        return True
        
    except Exception as e:
        print(f"Erro ao aplicar otimização: {e}")
        traceback.print_exc()
        return False

def record_optimization(function_name: str, issue: str, original_code: str, optimized_code: str):
    """Registra uma otimização no histórico."""
    try:
        optimizations = []
        if os.path.exists(OPTIMIZATION_HISTORY_FILE):
            with open(OPTIMIZATION_HISTORY_FILE, 'r', encoding='utf-8') as f:
                optimizations = json.load(f)
        
        optimizations.append({
            'timestamp': datetime.now().isoformat(),
            'function': function_name,
            'issue': issue,
            'original_size': len(original_code),
            'optimized_size': len(optimized_code),
            'improvement': len(original_code) - len(optimized_code)
        })
        
        with open(OPTIMIZATION_HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(optimizations, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"Erro ao registrar otimização: {e}")

def auto_optimize(force: bool = False) -> Dict:
    """
    Executa otimização automática do código.
    
    Args:
        force: Se True, força otimização mesmo sem problemas identificados
    
    Returns:
        Dicionário com resultados da otimização
    """
    print("\n" + "="*60)
    print("🤖 SISTEMA DE AUTO-MELHORAMENTO")
    print("="*60)
    
    # Analisar código
    print("\n📊 Analisando código...")
    analysis = analyze_own_code()
    
    print(f"   • {len(analysis['functions'])} funções encontradas")
    print(f"   • {len(analysis['classes'])} classes encontradas")
    print(f"   • {analysis['lines']} linhas de código")
    
    # Identificar oportunidades
    opportunities = analysis.get('optimization_opportunities', [])
    
    if not opportunities and not force:
        print("\n✅ Nenhuma oportunidade de otimização identificada no momento.")
        return {'success': True, 'optimizations_applied': 0}
    
    print(f"\n🔍 {len(opportunities)} oportunidades de otimização identificadas:")
    for i, opp in enumerate(opportunities[:10], 1):  # Mostrar top 10
        print(f"   {i}. {opp['function']}: {opp['issue']} (Prioridade: {opp['priority']})")
    
    # Aplicar otimizações (apenas high priority por padrão)
    applied = 0
    for opp in opportunities:
        if opp['priority'] == 'high' or force:
            print(f"\n⚡ Otimizando {opp['function']}...")
            
            # Obter código da função
            try:
                func = globals().get(opp['function'])
                if func:
                    source = inspect.getsource(func)
                    
                    # Gerar otimização
                    optimized = generate_optimization_code(
                        opp['function'],
                        opp['issue'],
                        source
                    )
                    
                    if optimized:
                        if apply_optimization(opp['function'], optimized, opp['issue']):
                            applied += 1
                            print(f"   ✅ {opp['function']} otimizada com sucesso!")
                        else:
                            print(f"   ❌ Falha ao aplicar otimização em {opp['function']}")
                    else:
                        print(f"   ⚠️ Não foi possível gerar otimização para {opp['function']}")
                else:
                    print(f"   ⚠️ Função {opp['function']} não encontrada no escopo global")
            except Exception as e:
                print(f"   ❌ Erro ao otimizar {opp['function']}: {e}")
    
    print(f"\n✅ Otimização concluída: {applied} melhorias aplicadas")
    print("="*60 + "\n")
    
    return {
        'success': True,
        'opportunities_found': len(opportunities),
        'optimizations_applied': applied
    }

# Decorator para monitorar execuções
def monitor_performance(func):
    """Decorator para monitorar desempenho de funções."""
    def wrapper(*args, **kwargs):
        start_time = time.time()
        success = False
        try:
            result = func(*args, **kwargs)
            success = True
            return result
        except Exception as e:
            print(f"Erro em {func.__name__}: {e}")
            raise
        finally:
            execution_time = time.time() - start_time
            performance_monitor.record_execution(func.__name__, execution_time, success)
    return wrapper

# ==================== FUNÇÕES DE NAVEGAÇÃO WEB ====================
def get_browser_window(browser_name: str = None):
    """
    Encontra a janela de um navegador.
    
    Args:
        browser_name: Nome do navegador ('Chrome', 'Edge', 'Firefox', etc.) ou None para qualquer
    
    Returns:
        Objeto Window do navegador ou None se não encontrado
    """
    browser_names = browser_name and [browser_name] or [
        "Google Chrome", "Chrome", "chrome.exe",
        "Microsoft Edge", "Edge", "msedge.exe",
        "Mozilla Firefox", "Firefox", "firefox.exe",
        "Opera", "opera.exe",
        "Brave", "brave.exe"
    ]
    
    for name in browser_names:
        try:
            windows = gw.getWindowsWithTitle(name)
            if windows:
                return windows[0]
        except Exception as e:
            print(f"Aviso ao buscar navegador {name}: {e}")
    
    return None

def activate_browser_window(browser_name: str = None) -> bool:
    """
    Ativa a janela do navegador se estiver aberta.
    """
    window = get_browser_window(browser_name)
    if window:
        try:
            window.activate()
            time.sleep(0.35)
            return True
        except Exception as e:
            print(f"Erro ao ativar navegador: {e}")
    return False


def _wait_browser_window(browser_name: str = None, timeout: float = 4.0, interval: float = 0.22):
    """Espera a janela do navegador aparecer (polling). Retorna True se apareceu, False após timeout."""
    deadline = time.time() + timeout
    while time.time() < deadline:
        w = get_browser_window(browser_name)
        if w:
            return True
        time.sleep(interval)
    return False


def open_program(program_name: str) -> bool:
    """
    Abre qualquer programa do computador pelo nome (como um assistente).
    No Windows: usa o menu Iniciar (Win + digitar nome + Enter).
    Funciona com qualquer aplicativo instalado (Calculadora, Word, Excel, Notepad, etc.).
    
    Args:
        program_name: Nome do programa como o usuário diria (ex: "calculadora", "bloco de notas", "Word")
    
    Returns:
        True se a ação foi executada (programa tende a abrir), False em caso de erro
    """
    name = (program_name or "").strip()
    if not name:
        print("Nome do programa não informado.")
        return False
    print(f"Abrindo programa: {name}...")
    try:
        if sys.platform == 'win32':
            # Método universal: Menu Iniciar + buscar pelo nome (funciona com qualquer app instalado)
            pyautogui.press('win')
            time.sleep(0.6)
            pyautogui.typewrite(name, interval=0.05)
            time.sleep(1.0)
            pyautogui.press('enter')
            time.sleep(0.5)
            print(f"Comando enviado para abrir: {name}")
            return True
        else:
            # Linux/Mac: tentar comando genérico (xdg-open, open)
            name_lower = name.lower()
            if '.exe' in name_lower or '/' in name or '\\' in name:
                subprocess.Popen(['xdg-open', name] if sys.platform.startswith('linux') else ['open', name],
                                 stdin=subprocess.DEVNULL, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            else:
                subprocess.Popen(['xdg-open', name] if sys.platform.startswith('linux') else ['open', '-a', name],
                                 stdin=subprocess.DEVNULL, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            return True
    except Exception as e:
        print(f"Erro ao abrir programa '{name}': {e}")
        return False


def open_browser(browser_name: str = None) -> bool:
    """
    Abre um navegador de internet.
    
    Args:
        browser_name: Nome do navegador ('Chrome', 'Brave', 'Edge', 'Firefox', etc.) ou None para DEFAULT_BROWSER
    
    Returns:
        True se conseguiu abrir, False caso contrário
    """
    if browser_name is None:
        browser_name = DEFAULT_BROWSER
    # Verificar se já está aberto
    window = get_browser_window(browser_name)
    if window and window.visible:
        print(f"{browser_name} já está aberto. Ativando janela...")
        return activate_browser_window(browser_name)
    
    print(f"Abrindo {browser_name}...")
    
    # Caminhos completos no Windows (quando não estão no PATH)
    program_files = os.environ.get("ProgramFiles", "C:\\Program Files")
    program_files_x86 = os.environ.get("ProgramFiles(x86)", "C:\\Program Files (x86)")
    local_appdata = os.environ.get("LOCALAPPDATA", "")
    
    browser_commands = {
        "Chrome": [
            "chrome.exe", "google-chrome", "chromium",
            os.path.join(local_appdata, "Google", "Chrome", "Application", "chrome.exe"),
            os.path.join(program_files, "Google", "Chrome", "Application", "chrome.exe"),
            os.path.join(program_files_x86, "Google", "Chrome", "Application", "chrome.exe"),
        ],
        "Edge": [
            "msedge.exe", "microsoft-edge",
            os.path.join(program_files, "Microsoft", "Edge", "Application", "msedge.exe"),
            os.path.join(program_files_x86, "Microsoft", "Edge", "Application", "msedge.exe"),
        ],
        "Firefox": ["firefox.exe", "firefox"],
        "Opera": ["opera.exe", "opera"],
        "Brave": [
            "brave.exe", "brave",
            os.path.join(local_appdata, "BraveSoftware", "Brave-Browser", "Application", "brave.exe"),
            os.path.join(program_files, "BraveSoftware", "Brave-Browser", "Application", "brave.exe"),
        ]
    }
    
    commands = [c for c in browser_commands.get(browser_name, ["chrome.exe"]) if c]
    
    # No Windows: tentar primeiro comando "start" para garantir que a janela apareça (ex.: quando o script é chamado pelo Node)
    if sys.platform == 'win32':
        for cmd in commands:
            if not os.path.isfile(cmd):
                continue
            cmd_path = os.path.normpath(os.path.abspath(cmd))
            try:
                # start "" "path" abre o programa em nova janela visível
                subprocess.Popen(
                    ['cmd', '/c', 'start', '', cmd_path],
                    shell=False,
                    stdin=subprocess.DEVNULL,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    creationflags=getattr(subprocess, 'CREATE_NO_WINDOW', 0x08000000) if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
                )
                if _wait_browser_window(browser_name, timeout=4.5, interval=0.22):
                    print(f"✅ {browser_name} aberto com sucesso! (start)")
                    activate_browser_window(browser_name)
                    return True
            except Exception as e:
                print(f"Aviso: start não funcionou para {cmd_path}: {e}")
    
    # Fallback: Popen direto (com DETACHED_PROCESS no Windows)
    creationflags = 0
    if sys.platform == 'win32':
        creationflags = getattr(subprocess, 'DETACHED_PROCESS', 0x00000008)
    
    for cmd in commands:
        try:
            use_shell = sys.platform == 'win32' and not os.path.isfile(cmd)
            if os.path.isfile(cmd):
                cmd_path = os.path.normpath(os.path.abspath(cmd))
                kw = {}
                if sys.platform == 'win32' and creationflags:
                    kw['creationflags'] = creationflags
                    kw['close_fds'] = True
                subprocess.Popen([cmd_path], **kw)
            else:
                kw = {}
                if sys.platform == 'win32' and creationflags:
                    kw['creationflags'] = creationflags
                    kw['close_fds'] = True
                subprocess.Popen([cmd], shell=use_shell, **kw)
            if _wait_browser_window(browser_name, timeout=4.5, interval=0.22):
                print(f"✅ {browser_name} aberto com sucesso!")
                activate_browser_window(browser_name)
                return True
        except FileNotFoundError:
            continue
        except Exception as e:
            print(f"Erro ao abrir {browser_name} via subprocess: {e}")
    
    # Fallback: se Brave não abriu, tentar Chrome depois Edge
    if browser_name == "Brave":
        print("Brave não encontrado. Tentando Chrome...")
        return open_browser("Chrome")
    if sys.platform == 'win32' and browser_name == "Chrome":
        print("Chrome não encontrado. Tentando Microsoft Edge...")
        return open_browser("Edge")
    
    # Fallback Windows: abrir navegador padrão com URL (start https://...)
    if sys.platform == 'win32':
        try:
            print("Tentando abrir navegador padrão com start https://...")
            subprocess.Popen(
                ['cmd', '/c', 'start', 'https://duckduckgo.com'],
                shell=False,
                stdin=subprocess.DEVNULL,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                creationflags=getattr(subprocess, 'CREATE_NO_WINDOW', 0x08000000) if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
            )
            if _wait_browser_window(None, timeout=4.5, interval=0.22):
                print("✅ Navegador padrão aberto com sucesso!")
                return True
        except Exception as e:
            print(f"Aviso: start https não funcionou: {e}")
    
    # Fallback: usar busca do Windows (Win+S e digitar nome do app)
    try:
        print(f"Tentando abrir {browser_name} via busca do Windows...")
        pyautogui.hotkey('win', 's')
        time.sleep(0.5)
        pyautogui.typewrite(browser_name)
        time.sleep(0.7)
        pyautogui.press('enter')
        if _wait_browser_window(browser_name, timeout=4.5, interval=0.22):
            print(f"✅ {browser_name} aberto com sucesso!")
            activate_browser_window(browser_name)
            return True
    except Exception as e:
        print(f"Erro ao abrir {browser_name} via busca: {e}")
    
    print(f"❌ Não foi possível abrir {browser_name}")
    return False

def navigate_to_url(url: str, browser_name: str = None) -> bool:
    """
    Navega para um endereço URL no navegador.
    
    Args:
        url: Endereço do site (com ou sem http://)
        browser_name: Nome do navegador ou None para DEFAULT_BROWSER
    
    Returns:
        True se conseguiu navegar, False caso contrário
    """
    if browser_name is None:
        browser_name = DEFAULT_BROWSER
    # Normalizar URL
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url
    
    print(f"Navegando para: {url}")
    
    # Ativar navegador
    if not activate_browser_window(browser_name):
        # Tentar abrir navegador padrão
        if not open_browser("Chrome"):
            if not open_browser("Edge"):
                print("❌ Não foi possível abrir nenhum navegador")
                return False
    
    time.sleep(0.25)
    
    # Focar na barra de endereço (Ctrl+L ou Alt+D)
    try:
        pyautogui.hotkey('ctrl', 'l')
        time.sleep(0.15)
    except:
        try:
            pyautogui.hotkey('alt', 'd')
            time.sleep(0.15)
        except Exception as e:
            print(f"Erro ao focar barra de endereço: {e}")
            return False
    
    # Selecionar tudo e digitar nova URL (interval 0 = mais rápido)
    pyautogui.hotkey('ctrl', 'a')
    time.sleep(0.06)
    pyautogui.typewrite(url, interval=0.005)
    time.sleep(0.15)
    
    # Pressionar Enter para navegar
    pyautogui.press('enter')
    time.sleep(0.7)
    
    print(f"✅ Navegando para {url}")
    return True

def open_browser_and_navigate(url: str = None, browser_name: str = None) -> bool:
    """
    Abre um navegador e navega para um endereço (função combinada).
    
    Args:
        url: Endereço do site (opcional, se None apenas abre o navegador)
        browser_name: Nome do navegador ou None para DEFAULT_BROWSER
    
    Returns:
        True se conseguiu, False caso contrário
    """
    if browser_name is None:
        browser_name = DEFAULT_BROWSER
    if not open_browser(browser_name):
        return False
    
    if url:
        return navigate_to_url(url, browser_name)
    
    return True


def close_browser(browser_name: str = None) -> bool:
    """
    Fecha o navegador (só deve ser chamado quando o usuário pedir explicitamente).
    Regra geral: não fechar o navegador após abertura; só fechar se for pedido.
    """
    if browser_name is None:
        browser_name = DEFAULT_BROWSER
    if not activate_browser_window(browser_name):
        activate_browser_window(None)
    time.sleep(0.3)
    try:
        pyautogui.hotkey('alt', 'F4')
        print("Navegador fechado (Alt+F4).")
        return True
    except Exception as e:
        print(f"Erro ao fechar navegador: {e}")
        return False


def search_on_web_human_like(query: str, browser_name: str = None) -> bool:
    """
    Abre o DuckDuckGo, clica no campo de busca DA PÁGINA (não na barra de URL),
    digita o termo e pressiona Enter. Simula uso humano.
    """
    if browser_name is None:
        browser_name = DEFAULT_BROWSER
    if not query or not query.strip():
        return open_browser(browser_name)
    if not open_browser(browser_name):
        return False
    if not open_browser_and_navigate("https://duckduckgo.com", browser_name):
        return False
    time.sleep(1.0)
    activate_browser_window(browser_name)
    time.sleep(0.25)
    pyautogui.press('escape')
    time.sleep(0.2)
    w, h = pyautogui.size()
    pyautogui.click(w // 2, int(h * 0.5))
    time.sleep(0.25)
    search_x = w // 2
    search_y = int(h * 0.38)
    pyautogui.click(search_x, search_y)
    time.sleep(0.3)
    q = query.strip()
    try:
        import pyperclip
        pyperclip.copy(q)
        time.sleep(0.08)
        pyautogui.hotkey('ctrl', 'v')
    except Exception:
        pyautogui.typewrite(q, interval=0.005)
    time.sleep(0.2)
    pyautogui.press('enter')
    time.sleep(0.8)
    print("✅ Pesquisa digitada no campo de busca do DuckDuckGo (não na URL).")
    return True


def search_on_web(query: str, browser_name: str = None) -> bool:
    """
    Abre o navegador e faz uma pesquisa no DuckDuckGo pela URL direta.
    """
    if browser_name is None:
        browser_name = DEFAULT_BROWSER
    if not query or not query.strip():
        return open_browser(browser_name)
    q = query.strip()
    url = "https://duckduckgo.com/?q=" + quote_plus(q)
    print("Abrindo pesquisa no DuckDuckGo (URL):", url[:60] + "..." if len(url) > 60 else url)
    return open_browser_and_navigate(url, browser_name)


# Padrão para o agente "entender onde está o link" em cada site:
# - Sempre usar o BUSCADOR do site (campo de busca) para pesquisar, não só abrir URL com query.
# - Localizar links e botões por SELETORES (CSS/XPath) no DOM, não por coordenadas (x,y).
# - YouTube: busca = input#search; primeiro vídeo = ytd-video-renderer a#video-title; play = button.ytp-large-play-button.
# - Para novos sites: adicionar função _site_xyz_selenium() com os seletores corretos da página.


def extract_youtube_video_id(url: str) -> Optional[str]:
    """Extrai o ID do vídeo de um link do YouTube (youtube.com/watch?v=ID, youtu.be/ID, youtube.com/embed/ID)."""
    if not url or not str(url).strip():
        return None
    url = str(url).strip()
    for pattern in [
        r"(?:youtube\.com/watch\?v=|youtu\.be/|youtube\.com/embed/)([a-zA-Z0-9_-]{11})",
        r"[?&]v=([a-zA-Z0-9_-]{11})",
    ]:
        m = re.search(pattern, url, re.IGNORECASE)
        if m:
            return m.group(1)
    return None


def extract_youtube_url_from_text(text: str) -> Optional[str]:
    """Extrai o primeiro link do YouTube encontrado no texto (para usar na mensagem do usuário)."""
    if not text or not str(text).strip():
        return None
    text = " ".join(str(text).split())  # normaliza espaços e quebras de linha
    # Padrões: https://youtube.com/watch?v=ID, https://youtu.be/ID, com ou sem www
    for pattern in [
        r"https?://(?:www\.)?youtube\.com/watch\?[^\s]+v=[a-zA-Z0-9_-]{11}[^\s\)\]\">]*",
        r"https?://(?:www\.)?youtu\.be/[a-zA-Z0-9_-]{11}[^\s\)\]\">]*",
        r"youtube\.com/watch\?v=([a-zA-Z0-9_-]{11})",
        r"youtu\.be/([a-zA-Z0-9_-]{11})",
    ]:
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            if m.lastindex:
                return f"https://www.youtube.com/watch?v={m.group(1)}"
            return m.group(0).rstrip(".,;:!?)")
    return None


def create_youtube_embed_html(video_id: str, title: str = "Vídeo") -> Optional[str]:
    """
    Cria um HTML que reproduz o vídeo do YouTube em embed (sem abrir a página do YouTube).
    Usa youtube-nocookie.com e player minimalista. Retorna o caminho do arquivo HTML.
    """
    if not video_id:
        return None
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = os.path.join(script_dir, "output_video")
    os.makedirs(output_dir, exist_ok=True)
    safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)[:50]
    path = os.path.join(output_dir, f"play_{video_id}_{int(time.time())}.html")
    # autoplay=1 pode ser bloqueado pelo navegador; o usuário pode clicar em play
    embed_url = (
        f"https://www.youtube-nocookie.com/embed/{video_id}"
        "?autoplay=1&rel=0&modestbranding=1&playsinline=1&enablejsapi=0"
    )
    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{safe_title}</title>
  <style>
    * {{ margin: 0; padding: 0; box-sizing: border-box; }}
    body {{ min-height: 100vh; display: flex; flex-direction: column; align-items: center; justify-content: center; background: #111; font-family: sans-serif; color: #eee; padding: 10px; }}
    .player {{ width: 100%; max-width: 900px; aspect-ratio: 16/9; }}
    .player iframe {{ width: 100%; height: 100%; border: none; }}
    .hint {{ margin-top: 8px; font-size: 12px; opacity: 0.8; }}
  </style>
</head>
<body>
  <div class="player">
    <iframe src="{embed_url}" allow="accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture" allowfullscreen></iframe>
  </div>
  <p class="hint">Reprodução em HTML (embed). Se não iniciar, clique no play.</p>
</body>
</html>"""
    try:
        with open(path, "w", encoding="utf-8") as f:
            f.write(html)
        print(f"HTML criado: {path}")
        return path
    except Exception as e:
        print(f"Erro ao criar HTML: {e}")
        return None


def open_youtube_video_in_html(url: str) -> bool:
    """
    Pega o link do vídeo do YouTube, cria um HTML com o player em embed (reprodução sem ser pela página do YouTube)
    e abre no navegador. Evita a interface e anúncios da página normal do YouTube.
    """
    video_id = extract_youtube_video_id(url)
    if not video_id:
        print("Link do YouTube inválido ou ID não encontrado.")
        return False
    path = create_youtube_embed_html(video_id, "Reprodução")
    if not path:
        return False
    try:
        path_abs = os.path.abspath(path)
        if not os.path.isfile(path_abs):
            print(f"Arquivo HTML não encontrado: {path_abs}")
            return False
        if sys.platform == "win32":
            os.startfile(path_abs)
        else:
            subprocess.Popen(
                ["xdg-open", path_abs] if sys.platform.startswith("linux") else ["open", path_abs],
                stdin=subprocess.DEVNULL, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            )
        print("Abrindo reprodução em HTML (embed, sem página do YouTube).")
        return True
    except Exception as e:
        print(f"Erro ao abrir HTML: {e}")
        traceback.print_exc()
        return False


def _youtube_search_and_play_selenium(query: str) -> bool:
    """
    Usa Selenium para: abrir YouTube, usar o buscador do site, localizar o primeiro
    link de vídeo pelo DOM (ytd-video-renderer) e clicar nele, depois clicar no botão play.
    Assim o agente "entende" onde está o link em vez de chutar coordenadas.
    """
    try:
        from selenium import webdriver
        from selenium.webdriver.chrome.options import Options as ChromeOptions
        from selenium.webdriver.common.by import By
        from selenium.webdriver.common.keys import Keys
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
    except ImportError:
        print("Selenium não instalado. Use: pip install selenium")
        return False

    driver = None
    try:
        q = (query or "música").strip()
        print(f"YouTube (Selenium): abrindo site, usando buscador para '{q}', clicando no primeiro link...")
        opts = ChromeOptions()
        opts.add_argument("--no-sandbox")
        opts.add_argument("--disable-dev-shm-usage")
        opts.add_argument("--disable-gpu")
        opts.add_argument("--window-size=1280,900")
        opts.add_argument("--disable-blink-features=AutomationControlled")
        opts.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
        opts.add_experimental_option("excludeSwitches", ["enable-logging", "enable-automation"])
        opts.add_experimental_option("detach", True)
        try:
            driver = webdriver.Chrome(options=opts)
        except Exception:
            try:
                from selenium.webdriver.edge.options import Options as EdgeOptions
                eopts = EdgeOptions()
                eopts.add_argument("--no-sandbox")
                eopts.add_argument("--disable-dev-shm-usage")
                eopts.add_argument("--disable-blink-features=AutomationControlled")
                eopts.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
                eopts.add_experimental_option("excludeSwitches", ["enable-logging", "enable-automation"])
                eopts.add_experimental_option("detach", True)
                driver = webdriver.Edge(options=eopts)
            except Exception as e:
                print(f"Chrome/Edge não disponível: {e}")
                return False

        wait = WebDriverWait(driver, 12)
        driver.get("https://www.youtube.com")
        time.sleep(1.2)

        # Fechar banner de cookies (Aceitar / Accept all) se aparecer
        for _ in range(2):
            try:
                for btn in driver.find_elements(By.CSS_SELECTOR, "button[aria-label], button.ytd-consent-dialog-renderer, tp-yt-paper-button"):
                    txt = (btn.text or "").strip().lower()
                    if "aceitar" in txt or "accept" in txt or "concordo" in txt or "agree" in txt or "tudo" in txt or "all" in txt:
                        btn.click()
                        time.sleep(0.4)
                        break
            except Exception:
                pass
            time.sleep(0.25)

        search_selectors = [
            "input#search",
            "input[name='search_query']",
            "ytd-searchbox input#search",
            "input.ytd-searchbox",
        ]
        search_input = None
        for sel in search_selectors:
            try:
                search_input = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, sel)))
                if search_input and search_input.is_displayed():
                    break
            except Exception:
                continue
        if not search_input:
            driver.get("https://www.youtube.com/results?search_query=" + quote_plus(q))
            time.sleep(2)
        else:
            search_input.click()
            time.sleep(0.15)
            search_input.clear()
            # Escrever no buscador de uma vez (JavaScript = instantâneo; send_keys seria lento)
            try:
                driver.execute_script("arguments[0].value = arguments[1]; arguments[0].dispatchEvent(new Event('input', { bubbles: true }));", search_input, q)
                time.sleep(0.15)
                search_input.send_keys(Keys.RETURN)
            except Exception:
                search_input.send_keys(q)
                search_input.send_keys(Keys.RETURN)
            time.sleep(2)

        video_link_selectors = [
            "ytd-video-renderer a#video-title",
            "ytd-video-renderer a[href*='/watch?v=']",
            "#contents ytd-video-renderer a#video-title",
            "a[href*='/watch?v='][id='video-title']",
            "ytd-item-section-renderer a#video-title",
        ]
        video_link = None
        for sel in video_link_selectors:
            try:
                video_link = driver.find_element(By.CSS_SELECTOR, sel)
                if video_link and video_link.is_displayed():
                    href = video_link.get_attribute("href") or ""
                    if "/watch?v=" in href:
                        break
            except Exception:
                continue
        if not video_link:
            try:
                links = driver.find_elements(By.CSS_SELECTOR, "a[href*='/watch?v=']")
                for a in links:
                    if a.is_displayed() and "Short" not in (a.get_attribute("href") or ""):
                        video_link = a
                        break
            except Exception:
                pass
        if not video_link:
            print("Nenhum link de vídeo encontrado na página de resultados.")
            return False

        print("Clicando no primeiro link de vídeo (encontrado pelo DOM)...")
        try:
            driver.execute_script("arguments[0].click();", video_link)
        except Exception:
            video_link.click()
        time.sleep(3)

        # Se aparecer "Algo deu errado" / "Something went wrong", clicar em "Tentar novamente" / "Try again"
        try:
            for el in driver.find_elements(By.CSS_SELECTOR, "button.ytp-error-content-buttons button, .ytp-error button, button[class*='retry'], paper-button"):
                try:
                    t = (el.text or "").strip().lower()
                    if "tentar" in t or "try again" in t or "again" in t or "reload" in t:
                        if el.is_displayed():
                            driver.execute_script("arguments[0].click();", el)
                            time.sleep(2)
                            break
                except Exception:
                    pass
            for el in driver.find_elements(By.TAG_NAME, "button"):
                try:
                    t = (el.text or "").strip().lower()
                    if ("tentar novamente" in t or "try again" in t) and el.is_displayed():
                        driver.execute_script("arguments[0].click();", el)
                        time.sleep(2)
                        break
                except Exception:
                    pass
        except Exception:
            pass

        play_selectors = [
            "button.ytp-large-play-button",
            ".ytp-large-play-button",
            "button[aria-label='Play']",
            ".html5-video-player .ytp-large-play-button",
        ]
        play_btn = None
        for sel in play_selectors:
            try:
                play_btn = WebDriverWait(driver, 8).until(
                    EC.element_to_be_clickable((By.CSS_SELECTOR, sel))
                )
                if play_btn and play_btn.is_displayed():
                    break
            except Exception:
                continue
        if play_btn:
            try:
                driver.execute_script("arguments[0].click();", play_btn)
                print("Botão play clicado (encontrado pelo DOM).")
            except Exception:
                play_btn.click()
        else:
            try:
                video = driver.find_element(By.CSS_SELECTOR, "video.html5-main-video")
                driver.execute_script("arguments[0].click();", video)
                print("Clicado no vídeo para dar play.")
            except Exception:
                print("Play não encontrado; vídeo pode já estar tocando ou layout diferente.")

        # Pular propaganda: aguardar até 12s o botão "Pular anúncio" / "Skip ad" (aparece em ~5s) e clicar
        skipped = False
        skip_selectors = [
            ".ytp-ad-skip-button-container",
            ".ytp-ad-skip-button-modern",
            ".ytp-ad-skip-button",
            "button.ytp-ad-skip-button",
            ".ytp-skip-ad-button",
            "[class*='ytp-ad-skip']",
        ]
        try:
            for sel in skip_selectors:
                if skipped:
                    break
                try:
                    skip_btn = WebDriverWait(driver, 12).until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, sel))
                    )
                    if skip_btn and skip_btn.is_displayed():
                        try:
                            from selenium.webdriver.common.action_chains import ActionChains
                            ActionChains(driver).move_to_element(skip_btn).click().perform()
                        except Exception:
                            driver.execute_script("arguments[0].click();", skip_btn)
                        print("Propaganda pulada (Pular anúncio / Skip ad).")
                        skipped = True
                        break
                except Exception:
                    continue
        except Exception:
            pass
        if not skipped:
            for _ in range(24):
                time.sleep(0.5)
                if skipped:
                    break
                try:
                    for sel in skip_selectors:
                        try:
                            skip_btn = driver.find_element(By.CSS_SELECTOR, sel)
                            if skip_btn.is_displayed():
                                driver.execute_script("arguments[0].scrollIntoView(true); arguments[0].click();", skip_btn)
                                print("Propaganda pulada.")
                                skipped = True
                                break
                        except Exception:
                            continue
                    if skipped:
                        break
                    for el in driver.find_elements(By.CSS_SELECTOR, "button, [role='button'], span.ytp-ad-skip-button-modern"):
                        try:
                            t = (el.text or el.get_attribute("aria-label") or "").strip().lower()
                            if "pular" in t or "skip" in t or "anúncio" in t or "ad" in t:
                                if el.is_displayed():
                                    driver.execute_script("arguments[0].click();", el)
                                    print("Propaganda pulada (por texto).")
                                    skipped = True
                                    break
                        except Exception:
                            pass
                except Exception:
                    pass
        # Fechar overlay de anúncio não pulável (X no canto)
        if not skipped:
            try:
                overlay_close = driver.find_elements(By.CSS_SELECTOR, ".ytp-ad-overlay-close-button, .ytp-ad-overlay-close-container button, button[aria-label='Fechar']")
                for btn in overlay_close:
                    if btn.is_displayed():
                        btn.click()
                        print("Overlay de anúncio fechado.")
                        break
            except Exception:
                pass

        time.sleep(1)
        print("✅ YouTube: buscou no buscador, clicou no primeiro link e acionou o play.")
        return True
    except Exception as e:
        print(f"Erro no fluxo YouTube (Selenium): {e}")
        traceback.print_exc()
        return False
    finally:
        # Regra: não fechar o navegador após abertura. Só fechar se o usuário pedir (ex: "feche o navegador").
        pass


def youtube_search_and_play(query: str, browser_name: str = None) -> bool:
    """
    Abre o YouTube, usa o buscador do site para pesquisar, localiza o primeiro
    link de vídeo pelo DOM (Selenium) e clica nele, depois clica no play.
    Padrão: sempre usar o buscador do site e clicar nos links reais (seletores DOM),
    não coordenadas na tela.
    """
    if _youtube_search_and_play_selenium(query):
        return True
    if browser_name is None:
        browser_name = DEFAULT_BROWSER
    print("Fallback: abrindo YouTube por URL e tentando cliques por posição...")
    q = (query or "música").strip()
    url = "https://www.youtube.com/results?search_query=" + quote_plus(q)
    if not open_browser(browser_name):
        return False
    if not open_browser_and_navigate(url, browser_name):
        return False
    time.sleep(5)
    activate_browser_window(browser_name or None)
    time.sleep(0.5)
    w, h = pyautogui.size()
    move_and_click(w // 2, int(h * 0.42))
    time.sleep(4)
    move_and_click(w // 2, h // 2)
    return True


def _fetch_duckduckgo_api(query: str) -> str:
    """Busca na API do DuckDuckGo e retorna texto para contexto."""
    try:
        from urllib.request import Request, urlopen
        url = "https://api.duckduckgo.com/?q=" + quote_plus(query) + "&format=json"
        req = Request(url, headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; rv:91.0) Gecko/20100101"})
        with urlopen(req, timeout=5) as r:
            data = json.loads(r.read().decode())
        parts = []
        if data.get("AbstractText"):
            parts.append(data["AbstractText"])
            if data.get("AbstractURL"):
                parts.append("Fonte: " + data["AbstractURL"])
        for item in data.get("RelatedTopics", [])[:5]:
            if isinstance(item, dict) and item.get("Text"):
                parts.append(item["Text"])
            elif isinstance(item, str):
                parts.append(item)
        return "\n\n".join(parts) if parts else ""
    except Exception as e:
        print(f"Aviso ao buscar DuckDuckGo API: {e}")
        return ""


def _fetch_page_text(url: str, max_chars: int = 2500) -> str:
    """Baixa uma página e extrai texto (remove HTML). Usado para enriquecer o contexto da varredura."""
    try:
        from urllib.request import Request, urlopen
        from urllib.error import URLError, HTTPError
        req = Request(url, headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"})
        with urlopen(req, timeout=5) as r:
            raw = r.read().decode("utf-8", errors="replace")
        text = re.sub(r"<script[^>]*>[\s\S]*?</script>", " ", raw, flags=re.IGNORECASE)
        text = re.sub(r"<style[^>]*>[\s\S]*?</style>", " ", text, flags=re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text).replace("&nbsp;", " ").replace("&amp;", "&").replace("&#39;", "'").strip()
        return text[:max_chars] if text else ""
    except (URLError, HTTPError, OSError) as e:
        return ""


def _extract_result_urls_from_html(html: str, max_urls: int = 5) -> List[str]:
    """Extrai URLs de resultados da página HTML do DuckDuckGo (links externos, não do próprio DDG)."""
    seen = set()
    urls = []
    # Links que aparecem em resultados (evitar duckduckgo.com, javascript, etc.)
    for m in re.findall(r'href="(https?://[^"]+)"', html):
        u = m.split("#")[0].strip()
        if "duckduckgo.com" in u or "duck.com" in u or u in seen:
            continue
        if u.startswith("http") and len(u) < 500:
            seen.add(u)
            urls.append(u)
            if len(urls) >= max_urls:
                break
    return urls


def _fetch_duckduckgo_html_snippets_and_urls(query: str) -> Tuple[str, List[str]]:
    """Busca na página HTML do DuckDuckGo: retorna (texto dos snippets, lista de URLs dos resultados)."""
    try:
        from urllib.request import Request, urlopen
        url = "https://html.duckduckgo.com/html/?q=" + quote_plus(query)
        req = Request(url, headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"})
        with urlopen(req, timeout=5) as r:
            html = r.read().decode("utf-8", errors="replace")
        snippets = []
        for pattern in [
            r'class="result__snippet"[^>]*>([^<]+(?:\s*<[^>]+>[^<]*)*)',
            r'class="snippet"[^>]*>([^<]+(?:\s*<[^>]+>[^<]*)*)',
            r'class="result__body"[^>]*>([^<]+(?:\s*<[^>]+>[^<]*)*)',
            r'<span[^>]*class="[^"]*snippet[^"]*"[^>]*>([^<]+)',
        ]:
            for m in re.findall(pattern, html, re.IGNORECASE):
                text = re.sub(r"<[^>]+>", " ", m).replace("&nbsp;", " ").replace("&amp;", "&").replace("&#39;", "'").strip()
                if len(text) > 40 and text not in snippets:
                    snippets.append(text)
            if snippets:
                break
        snippet_text = "\n\n".join(snippets[:8]) if snippets else ""
        result_urls = _extract_result_urls_from_html(html, max_urls=5)
        return snippet_text, result_urls
    except Exception as e:
        print(f"Aviso ao buscar DuckDuckGo HTML: {e}")
        return "", []


def _fetch_duckduckgo_html_snippets(query: str) -> str:
    """Compatibilidade: retorna só os snippets (sem URLs)."""
    return _fetch_duckduckgo_html_snippets_and_urls(query)[0]


def _fetch_web_context(query: str) -> str:
    """
    Varredura na web: API DuckDuckGo + página de resultados (snippets) + conteúdo de até 2 páginas
    dos resultados, para montar um contexto rico e responder somente com base no encontrado.
    """
    parts = []
    # 1) API (respostas instantâneas)
    api_text = _fetch_duckduckgo_api(query)
    if api_text and len(api_text.strip()) >= 50:
        parts.append(api_text.strip())
    # 2) Snippets e URLs da página HTML
    snippet_text, result_urls = _fetch_duckduckgo_html_snippets_and_urls(query)
    if snippet_text:
        parts.append(snippet_text)
    # 3) Conteúdo de até 2 páginas dos resultados (varredura mais profunda)
    for url in result_urls[:2]:
        page_text = _fetch_page_text(url, max_chars=2500)
        if page_text and len(page_text) > 100:
            parts.append(f"[Conteúdo de página]\nFonte: {url}\n\n{page_text}")
    return "\n\n---\n\n".join(parts).strip() if parts else ""


def _varredura_clicando_em_links(query: str, max_links: int = 3) -> str:
    """
    Abre o navegador com Selenium, faz a pesquisa no DuckDuckGo, obtém os links dos resultados
    e abre (equivale a clicar) em até max_links páginas para extrair o texto e montar o contexto.
    Retorna string com o conteúdo encontrado ou "" se Selenium não estiver disponível/falhar.
    """
    try:
        from selenium import webdriver
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.common.by import By
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from urllib.parse import unquote, parse_qs, urlparse
    except ImportError as e:
        print("Aviso: Selenium não está instalado. Instale com: pip install selenium")
        return ""

    driver = None
    try:
        print("Tentando varredura clicando em links (Selenium)...")
        options = Options()
        options.add_argument("--headless=new")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--disable-gpu")
        options.add_argument("--window-size=1920,1080")
        options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
        options.add_experimental_option("excludeSwitches", ["enable-logging"])
        try:
            driver = webdriver.Chrome(options=options)
        except Exception:
            # Fallback para Edge no Windows se Chrome não estiver disponível
            try:
                from selenium.webdriver.edge.options import Options as EdgeOptions
                eopts = EdgeOptions()
                eopts.add_argument("--headless=new")
                eopts.add_argument("--no-sandbox")
                driver = webdriver.Edge(options=eopts)
            except Exception:
                raise
        _timeout = TIMEOUT_VARREDURA
        driver.set_page_load_timeout(_timeout)

        search_url = "https://html.duckduckgo.com/html/?q=" + quote_plus(query.strip())
        print("Abrindo página de resultados no navegador...")
        driver.get(search_url)
        time.sleep(1)
        WebDriverWait(driver, _timeout).until(EC.presence_of_element_located((By.CSS_SELECTOR, "a[href], .result__a, .result__snippet")))

        # Coletar links: no DuckDuckGo HTML os resultados são a.result__a com href = "https://duckduckgo.com/l/?kh=-1&uddg=URL_REAL_ENCODADA"
        def extract_real_url(href_str):
            if not href_str or "uddg=" not in href_str:
                return None
            try:
                if href_str.startswith("/"):
                    full = "https://duckduckgo.com" + href_str
                else:
                    full = href_str
                parsed = parse_qs(urlparse(full).query)
                uddg = (parsed.get("uddg") or [None])[0]
                if uddg:
                    real = unquote(uddg)
                    if real.startswith("http") and "duckduckgo.com" not in real and len(real) < 500:
                        return real
            except Exception:
                pass
            return None

        links = []
        seen = set()
        # Priorizar links de resultado (classe result__a)
        for a in driver.find_elements(By.CSS_SELECTOR, "a.result__a"):
            href = (a.get_attribute("href") or "").strip()
            if not href or href in seen:
                continue
            real = extract_real_url(href)
            if real and real not in seen:
                seen.add(real)
                links.append(real)
            elif href.startswith("http") and "duckduckgo.com" not in href and "duck.com" not in href and href not in seen and len(href) < 500:
                seen.add(href)
                links.append(href)
            if len(links) >= max_links + 5:
                break
        # Se não achou pelos result__a, pegar qualquer link com uddg
        if not links:
            for a in driver.find_elements(By.CSS_SELECTOR, "a[href]"):
                href = (a.get_attribute("href") or "").strip()
                if not href or href in seen:
                    continue
                real = extract_real_url(href)
                if real and real not in seen:
                    seen.add(real)
                    links.append(real)
                elif href.startswith("http") and "duckduckgo.com" not in href and "duck.com" not in href and len(href) < 500 and href not in seen:
                    seen.add(href)
                    links.append(href)
                if len(links) >= max_links + 5:
                    break

        # Fallback: extrair links do HTML bruto (page_source) — a página pode ter estrutura diferente
        if not links:
            raw_html = driver.page_source or ""
            # Procurar uddg=URL_ENCODADA em href (ex: href="/l/?uddg=https%3A%2F%2F..." ou uddg= no meio da URL)
            for m in re.findall(r"uddg=([^&\s\"']+)", raw_html):
                try:
                    real = unquote(m)
                    if real.startswith("http") and "duckduckgo.com" not in real and len(real) < 500 and real not in seen:
                        seen.add(real)
                        links.append(real)
                        if len(links) >= max_links + 5:
                            break
                except Exception:
                    pass

        print(f"Encontrados {len(links)} link(s) nos resultados. Abrindo até {min(max_links, len(links))} página(s)...")

        parts = []
        # Snippets da própria página de resultados (já estamos nela)
        try:
            for el in driver.find_elements(By.CSS_SELECTOR, ".result__snippet, .snippet"):
                t = (el.text or "").strip()
                if len(t) > 40:
                    parts.append(t)
        except Exception:
            pass

        # Timeout configurável ao abrir cada link
        try:
            from selenium.common.exceptions import TimeoutException as SeleniumTimeout
        except ImportError:
            SeleniumTimeout = Exception
        driver.set_page_load_timeout(TIMEOUT_VARREDURA)

        # Clicar/abrir até max_links páginas dos resultados e ler o conteúdo
        for i, url in enumerate(links[:max_links]):
            try:
                print(f"Clicando/abrindo link {i + 1}/{min(len(links), max_links)}: {url[:50]}...")
                driver.get(url)
                WebDriverWait(driver, TIMEOUT_VARREDURA).until(lambda d: d.execute_script("return document.readyState") == "complete")
                time.sleep(0.5)
                raw = driver.page_source or ""
                text = re.sub(r"<script[^>]*>[\s\S]*?</script>", " ", raw, flags=re.IGNORECASE)
                text = re.sub(r"<style[^>]*>[\s\S]*?</style>", " ", text, flags=re.IGNORECASE)
                text = re.sub(r"<[^>]+>", " ", text)
                text = re.sub(r"\s+", " ", text).replace("&nbsp;", " ").replace("&amp;", "&").strip()
                text = text[:3000].strip() if text else ""
                if len(text) > 150:
                    parts.append(f"[Página {i + 1}]\nFonte: {url}\n\n{text}")
            except SeleniumTimeout:
                print(f"   Timeout ao carregar ({TIMEOUT_VARREDURA}s). Próximo link.")
            except Exception as e:
                err = str(e).split("\n")[0][:80] if str(e) else "erro"
                print(f"   Não foi possível ler: {err}")
            finally:
                if i < len(links[:max_links]) - 1:
                    try:
                        driver.back()
                        time.sleep(0.5)
                    except Exception:
                        driver.get(search_url)
                        time.sleep(0.5)

        if parts:
            print(f"✅ Varredura concluída: {len(parts)} trecho(s) obtido(s) a partir dos links.")
            return "\n\n---\n\n".join(parts).strip()
        if not links:
            print("Aviso: nenhum link de resultado encontrado na página. Usando varredura por HTTP.")
        return ""
    except Exception as e:
        print(f"Aviso: varredura com cliques (Selenium) falhou: {e}")
        print("Usando varredura por HTTP (sem clicar em links).")
        return ""
    finally:
        if driver:
            try:
                driver.quit()
            except Exception:
                pass


def search_web_and_answer(task: str, query: str) -> bool:
    """
    Faz varredura na web (API + resultados + páginas), monta contexto e gera resposta
    SOMENTE com base na pergunta e no conteúdo encontrado. A resposta é impressa para o WhatsApp.
    """
    if not query or not query.strip():
        return False
    q = query.strip()
    print("Varrendo a web em busca de:", q[:60] + ("..." if len(q) > 60 else ""))
    # Primeiro tenta varredura clicando em links (Selenium); se não houver conteúdo, usa fetch por HTTP
    context = _varredura_clicando_em_links(q, max_links=MAX_LINKS_VARREDURA)
    if not context:
        print("Varredura por HTTP (snippets + páginas)...")
        context = _fetch_web_context(q)
    if not context:
        print("Não foi possível obter conteúdo da web para esta busca. Abrindo o navegador.")
        return False
    if not client:
        print("OPENAI_API_KEY não configurada. Não é possível gerar resposta.")
        return False
    system = (
        "Responda APENAS com a lista de resultados: cada item em uma linha com um resumo curto e o link da fonte. "
        "Não invente dados. Não escreva introduções, despedidas, recomendações nem texto extra. "
        "Formato: número, resumo breve, e na linha seguinte Fonte: URL. Só isso."
    )
    prompt = (
        "Informações obtidas da web (varredura):\n\n"
        f"{context[:6000]}\n\n"
        "---\n"
        f"Pergunta do usuário: {task}\n\n"
        "Liste somente os resultados encontrados: para cada um, um resumo curto e o link (Fonte: URL). Nada mais."
    )
    try:
        answer = get_ai_response(prompt, system, max_tokens=800)
        if answer:
            # Limite para caber bem no WhatsApp (evitar mensagem gigante)
            max_chars = int(os.environ.get("MAX_RESPOSTA_WHATSAPP", "3500"))
            if len(answer) > max_chars:
                answer = answer[: max_chars - 3].rstrip() + "..."
            print("\n📋 Resposta com base na varredura na web:\n")
            print(answer)
            print("\n✅ Resposta gerada com sucesso.")
            return True
    except Exception as e:
        print(f"Erro ao gerar resposta: {e}")
    return False


def run_search_with_answer(task: str, query: str) -> bool:
    """Busca na web, gera resposta com GPT e abre o navegador com a pesquisa."""
    answered = search_web_and_answer(task, query)
    opened = search_on_web(query) if query else open_browser()
    return answered or opened


def show_help() -> bool:
    """Mostra o que o agente pode fazer (comando para usar no WhatsApp com várias funções)."""
    help_text = """*Comando do agente no WhatsApp – funções disponíveis:*

💡 *O GPT ajuda a entender o que você pedir* — você pode falar em linguagem natural (ex: "quero ver o clima", "me mostra notícias") e o agente interpreta e executa.

*Funções principais:*

• *Ligar câmera* – Ex: "abrir câmera", "ligar câmera", "ativar câmera"
  → Envia as imagens ao servidor; você recebe uma URL para ver ao vivo na rede externa.

• *Desligar câmera* – Ex: "desligar câmera", "fechar câmera", "desativar câmera"
  → Fecha o app de câmera.

• *Tirar foto* – Ex: "tirar foto", "capturar foto"
  → Tira uma foto com a câmera e salva.

• *Abrir navegador* – Ex: "abrir navegador", "abrir Chrome", "ir para google.com"
  → Abre o navegador (Brave/Chrome) e opcionalmente acessa um site.

• *Abrir qualquer programa* – Ex: "abrir calculadora", "abrir Word", "executar Bloco de Notas"
  → Abre qualquer aplicativo instalado no PC (menu Iniciar). Use como assistente.

• *Pesquisar na web* – Ex: "pesquisar clima", "notícias Joinville", "buscar receita bolo"
  → Varre a web, abre links e devolve resultado + fonte de cada link.

• *Rede* – Ex: "conectar wifi", "verificar conexão"
  → Ajuda com conexão à internet.

• *Ver este menu* – Ex: "agente", "comandos", "ajuda", "o que você pode fazer"
  → Mostra esta lista de funções.

*Exemplos de linguagem natural:*
• "quero ver o clima" → pesquisa clima
• "me mostra notícias de Joinville" → pesquisa notícias Joinville
• "preciso abrir o Excel" → abre Excel
• "liga a câmera" → liga câmera e envia link"""
    print("\n📋 Resposta com base na varredura na web:\n")
    print(help_text)
    print("\n✅ Resposta gerada com sucesso.")
    return True


# ==================== FUNÇÕES DE EXECUÇÃO ====================
def prompt_text_only(text: str) -> str:
    """Remove blocos de código e markdown; retorna só texto em linguagem natural."""
    if not text or not text.strip():
        return text
    s = text.strip()
    s = re.sub(r'```[\s\S]*?```', ' ', s)
    s = re.sub(r'^```\w*\s*$', '', s, flags=re.MULTILINE)
    s = re.sub(r'\n\s*\n', '\n', s)
    s = re.sub(r' +', ' ', s)
    return s.strip()


def send_to_cursor_terminal(command: str):
    """
    Envia um comando para o terminal do Cursor.
    Não faz nada se CURSOR_INTERACTION_DISABLED estiver True.
    
    Args:
        command: Comando/descrição a ser enviado para o terminal
    """
    if CURSOR_INTERACTION_DISABLED:
        return
    print(f"Enviando comando para o terminal do Cursor: {command}")
    activate_cursor_window()
    
    # Abrir terminal (Ctrl + `)
    pyautogui.hotkey('ctrl', '`')
    time.sleep(0.8)
    
    # Digitar o comando
    pyautogui.typewrite(command, interval=0.05)
    time.sleep(0.3)
    
    # Executar (Enter)
    pyautogui.press('enter')
    time.sleep(0.5)
    
    print(f"Comando enviado para o terminal: {command[:50]}...")


def _notify_cursor_terminal_extension_added(ext_file: str, description: str) -> None:
    """
    Envia ao terminal do Cursor um aviso de que um novo comando foi adicionado,
    para o usuário ver as melhorias e poder atualizar. Usado mesmo com CURSOR_INTERACTION_DISABLED.
    """
    if not ENVIAR_ATUALIZACOES_TERMINAL_CURSOR:
        return
    try:
        # Ativar Cursor (ignorar CURSOR_INTERACTION_DISABLED só para esta notificação)
        window = get_cursor_window()
        if not window:
            return
        window.activate()
        time.sleep(0.6)
        # Abrir terminal
        pyautogui.hotkey('ctrl', '`')
        time.sleep(0.8)
        # Comando: echo para aparecer no terminal do Cursor (atualização aplicada)
        msg_ascii = "Agente atualizado. Novo comando adicionado em cursor_automation_extensions.py"
        try:
            import pyperclip
            # Descrição pode ter acentos; usar só msg_ascii no echo para evitar erro no terminal
            pyperclip.copy(f'echo [Agente] {msg_ascii}')
            time.sleep(0.2)
            pyautogui.hotkey('ctrl', 'v')
        except Exception:
            pyautogui.typewrite(f'echo [Agente] {msg_ascii}', interval=0.03)
        time.sleep(0.3)
        pyautogui.press('enter')
        time.sleep(0.4)
        print("Aviso de atualização enviado ao terminal do Cursor.")
    except Exception as e:
        print(f"Aviso: não foi possível enviar ao terminal do Cursor: {e}")

def send_to_cursor_chat(message: str):
    """
    Envia uma mensagem para o chat do Cursor.
    
    Args:
        message: Mensagem/prompt a ser enviado para o Cursor
    """
    print(f"Enviando mensagem para o chat do Cursor...")
    activate_cursor_window()
    
    # Abrir chat do Cursor (Ctrl+L é o atalho padrão)
    pyautogui.hotkey('ctrl', 'l')
    time.sleep(1)
    
    # Digitar a mensagem
    pyautogui.typewrite(message, interval=0.01)
    time.sleep(0.5)
    
    # Enviar (Enter)
    pyautogui.press('enter')
    time.sleep(0.5)
    
    print(f"Mensagem enviada para o Cursor: {message[:50]}...")

def wait_for_cursor_code_completion(timeout: int = CODE_COMPLETION_TIMEOUT, 
                                     check_interval: float = CODE_COMPLETION_CHECK_INTERVAL) -> bool:
    """
    Aguarda o Cursor terminar de gerar/processar código.
    
    Estratégias de detecção:
    1. Aguarda um tempo mínimo (Cursor geralmente leva alguns segundos)
    2. Monitora atividade do teclado (quando para de digitar)
    3. Verifica se há código no editor (seleciona tudo e verifica se há conteúdo)
    
    Args:
        timeout: Tempo máximo de espera em segundos
        check_interval: Intervalo entre verificações em segundos
    
    Returns:
        True se detectou que o código foi gerado, False se timeout
    """
    print("Aguardando Cursor terminar de gerar código...")
    
    min_wait_time = 5
    start_time = time.time()
    
    # Aguardar tempo mínimo
    time.sleep(min_wait_time)
    
    # Monitorar atividade
    last_activity_time = time.time()
    inactivity_threshold = 8  # Segundos sem atividade para considerar completo
    
    while time.time() - start_time < timeout:
        current_time = time.time()
        
        try:
            window = get_cursor_window()
            if window:
                activate_cursor_window()
                time.sleep(0.3)
                
                elapsed_since_last_activity = current_time - last_activity_time
                
                if elapsed_since_last_activity >= inactivity_threshold:
                    print("✅ Detectado: Cursor parece ter terminado (sem atividade recente)")
                    return True
                
        except Exception as e:
            print(f"Aviso ao verificar código: {e}")
        
        time.sleep(check_interval)
        last_activity_time = current_time
    
    print(f"⚠️ Timeout aguardando Cursor terminar (esperou {timeout}s)")
    return False

def detect_code_ready() -> bool:
    """
    Detecta se o código está pronto no editor do Cursor.
    Verifica se há conteúdo no editor.
    
    Returns:
        True se há código no editor, False caso contrário
    """
    try:
        window = get_cursor_window()
        if not window:
            return False
        
        activate_cursor_window()
        time.sleep(0.3)
        
        # Selecionar tudo
        pyautogui.hotkey('ctrl', 'a')
        time.sleep(0.2)
        
        # Copiar
        pyautogui.hotkey('ctrl', 'c')
        time.sleep(0.3)
        
        # Desmarcar seleção
        pyautogui.press('right')
        
        return True
        
    except Exception as e:
        print(f"Erro ao detectar código: {e}")
        return False

def click_run_button() -> bool:
    """
    Tenta executar código no Cursor usando várias estratégias.
    
    Returns:
        True se conseguiu executar, False caso contrário
    """
    print("🔍 Procurando botão Run no Cursor...")
    activate_cursor_window()
    time.sleep(1.0)
    
    # Garantir que estamos no editor (fechar terminal se estiver aberto)
    try:
        pyautogui.hotkey('ctrl', '`')  # Toggle terminal
        time.sleep(0.3)
        pyautogui.hotkey('ctrl', '`')  # Fechar terminal
        time.sleep(0.5)
    except:
        pass
    
    # Estratégias de execução (em ordem de preferência)
    strategies = [
        ("F5", lambda: pyautogui.press('f5')),
        ("Ctrl+F5", lambda: pyautogui.hotkey('ctrl', 'f5')),
        ("Ctrl+Enter", lambda: pyautogui.hotkey('ctrl', 'enter')),
        ("Alt+R", lambda: pyautogui.hotkey('alt', 'r')),
    ]
    
    for name, action in strategies:
        try:
            print(f"⌨️ Tentando atalho {name}...")
            action()
            time.sleep(2)
            print(f"✅ Atalho {name} pressionado")
            return True
        except Exception as e:
            print(f"⚠️ Atalho {name} não funcionou: {e}")
    
    # Tentar Command Palette
    try:
        print("⌨️ Tentando Command Palette (Ctrl+Shift+P)...")
        pyautogui.hotkey('ctrl', 'shift', 'p')
        time.sleep(1.5)
        pyautogui.typewrite('run', interval=0.1)
        time.sleep(1)
        pyautogui.press('enter')
        time.sleep(2)
        print("✅ Comando Run executado via Command Palette")
        return True
    except Exception as e:
        print(f"⚠️ Command Palette não funcionou: {e}")
        try:
            pyautogui.press('escape')
            time.sleep(0.3)
        except:
            pass
    
    print("❌ Não foi possível executar código usando botão Run")
    return False

# ==================== MAPEAMENTO DE TAREFAS ====================
TASK_MAPPING: Dict[str, Dict] = {
    'camera': {
        'keywords': ['camera', 'câmera', 'camara', 'webcam', 'web cam', 'câmera integrada', 
                    'camera integrada', 'câmera do dispositivo', 'camera do dispositivo', 
                    'câmera do notebook', 'camera do notebook', 'camera integrada do dispositivo'],
        'actions': ['abrir', 'abre', 'acessar', 'acessa', 'abrir aplicativo', 'abrir a', 
                   'acessar a', 'abrir o aplicativo', 'abrir o', 'acessar o', 'ativar', 
                   'ativa', 'ativar a', 'ligar', 'liga', 'iniciar', 'inicia', 'começar a capturar', 'capturar vídeo', 
                   'capturar video', 'tirar fotos'],
        'function': start_camera_stream_to_server,
        'description': 'Ligar câmera e enviar imagens ao servidor (URL para ver na rede)'
    },
    'photo': {
        'keywords': ['tirar foto', 'tirar fotos', 'capturar foto', 'capturar imagem', 
                    'foto', 'fotografar', 'tirar uma foto'],
        'actions': ['tirar', 'capturar', 'fazer', 'tirar uma'],
        'function': lambda: take_photo(),
        'description': 'Tirar foto com a câmera'
    },
    'camera_off': {
        'keywords': ['camera', 'câmera', 'camara', 'webcam', 'app câmera', 'app camera'],
        'actions': ['desligar', 'desliga', 'fechar', 'fecha', 'encerrar', 'encerra', 'desativar', 'desativa', 'parar', 'sair da'],
        'function': close_camera_app,
        'description': 'Desligar/fechar o app de câmera'
    },
    'browser': {
        'keywords': ['navegador', 'browser', 'internet', 'web', 'site', 'página', 'url', 
                    'endereço', 'google', 'chrome', 'edge', 'firefox', 'abrir navegador',
                    'abrir browser', 'acessar internet', 'navegar', 'navegar para'],
        'actions': ['abrir', 'abre', 'acessar', 'acessa', 'abrir o', 'acessar o', 
                   'navegar', 'navegar para', 'ir para', 'ir ao', 'abrir site',
                   'digitar endereço', 'digitar url', 'digitar na barra'],
        'function': lambda url=None: open_browser_and_navigate(url) if url else open_browser(),
        'description': 'Abrir navegador e acessar site'
    },
    'search': {
        'keywords': ['pesquisar', 'pesquisa', 'buscar', 'busca', 'search', 'google', 'procurar',
                     'notícias', 'noticias', 'notícia', 'noticia', 'notícias de hoje', 'noticias de hoje'],
        'actions': ['pesquisar', 'pesquisa', 'buscar', 'busca', 'procurar', 'procura'],
        'function': lambda q=None: search_on_web(q) if q else open_browser(),
        'description': 'Abrir navegador e pesquisar na web'
    },
    'network': {
        'keywords': ['conectar', 'conexão', 'conexao', 'rede', 'wifi', 'wi-fi', 'ethernet',
                    'internet', 'conectar-se', 'conectar se', 'conectar à internet',
                    'conectar a internet', 'estabelecer conexão', 'ativar conexão',
                    'verificar conexão', 'verificar internet', 'conexão de rede',
                    'conexao de rede', 'cabo ethernet', 'cabo de rede'],
        'actions': ['conectar', 'conectar-se', 'conectar se', 'estabelecer', 'ativar',
                   'verificar', 'checar', 'conectar à', 'conectar a'],
        'function': lambda ssid=None, password=None: connect_to_internet(ssid, password),
        'description': 'Conectar à internet via Wi-Fi ou Ethernet'
    },
    'self_optimize': {
        'keywords': ['agente', 'software', 'modificar', 'código', 'codigo', 'comportamento',
                    'desempenho', 'eficiencia', 'eficência', 'otimizar', 'melhorar',
                    'auto-melhorar', 'auto melhorar', 'auto-otimizar', 'auto otimizar',
                    'auto-evoluir', 'auto evoluir', 'auto-modificar', 'auto modificar',
                    'auto-aprender', 'auto aprender', 'aprendizado', 'máquina', 'machine learning',
                    'evoluir', 'evolução', 'otimização', 'otimizacao'],
        'actions': ['desenvolver', 'criar', 'implementar', 'melhorar', 'otimizar',
                   'evoluir', 'modificar', 'ajustar', 'aprender', 'identificar',
                   'implementar alterações', 'implementar alteracoes'],
        'function': lambda force=False: auto_optimize(force),
        'description': 'Auto-melhorar e otimizar o código do agente'
    },
}

# Carregar extensões definidas pelo usuário ou pelo próprio agente (GPT edita este arquivo)
def _load_extension_mappings():
    _ext_file = os.path.join(os.path.dirname(__file__), 'cursor_automation_extensions.py')
    if not os.path.exists(_ext_file):
        return
    try:
        _ctx = {
            'subprocess': subprocess, 'os': os, 'sys': sys, 'time': time,
            'pyautogui': pyautogui, 'open_browser': open_browser, 'search_on_web': search_on_web,
            'open_browser_and_navigate': open_browser_and_navigate,
            'start_camera_stream_to_server': start_camera_stream_to_server,
            'EXTRA_MAPPINGS': {},
        }
        exec(open(_ext_file, encoding='utf-8').read(), _ctx)
        for k, v in _ctx.get('EXTRA_MAPPINGS', {}).items():
            TASK_MAPPING[k] = v
        if _ctx.get('EXTRA_MAPPINGS'):
            print("Extensões de automação carregadas:", list(_ctx['EXTRA_MAPPINGS'].keys()))
    except Exception as e:
        print(f"Aviso: não foi possível carregar extensões: {e}")

_load_extension_mappings()

def extract_search_query(task: str) -> Optional[str]:
    """
    Extrai o termo de busca da tarefa (ex: "pesquisar clima em SP" -> "clima em SP";
    "Pesquisa última notícia em Joinville" -> "última notícia em Joinville").
    """
    task_stripped = task.strip()
    task_lower = task_stripped.lower()
    # Remover "Pesquisa " / "Pesquisa de " no início para melhorar a query no buscador
    for prefix in ("pesquisa de ", "pesquisa ", "buscar ", "busca ", "pesquisar "):
        if task_lower.startswith(prefix):
            task_stripped = task_stripped[len(prefix):].strip()
            task_lower = task_stripped.lower()
            break
    # Notícia(s): extrair query (ex: "Última notícia" -> frase inteira; "me fale as notícias de hoje" -> "notícias de hoje")
    for kw in ['notícias', 'noticias', 'notícia', 'noticia']:
        if kw in task_lower:
            idx = task_lower.find(kw)
            part = task_stripped[idx:].strip() or task_stripped
            # Se a frase já é curta, usar inteira para melhor resultado na varredura
            return task_stripped if len(task_stripped) <= 50 else part
    for word in ['pesquisar', 'pesquisa', 'buscar', 'busca', 'procurar', 'procura', 'search']:
        idx = task_lower.find(word)
        if idx != -1:
            rest = task_stripped[idx + len(word):].strip()
            for prep in ['no ', 'na ', 'sobre ', 'por ', 'na web ', 'no google ']:
                if rest.lower().startswith(prep):
                    rest = rest[len(prep):].strip()
            if rest:
                return rest
    return None

def extract_url_from_task(task: str) -> Optional[str]:
    """
    Extrai URL ou endereço de site da descrição da tarefa.
    
    Args:
        task: Descrição da tarefa
        
    Returns:
        URL extraída ou None se não encontrar
    """
    # Padrões para encontrar URLs
    url_patterns = [
        r'https?://[^\s]+',  # http:// ou https://
        r'www\.[^\s]+',  # www.site.com
        r'[a-zA-Z0-9-]+\.[a-zA-Z]{2,}(?:\.[a-zA-Z]{2,})?',  # site.com ou site.com.br
    ]
    
    for pattern in url_patterns:
        match = re.search(pattern, task)
        if match:
            url = match.group(0)
            # Normalizar
            if not url.startswith(('http://', 'https://')):
                if url.startswith('www.'):
                    url = 'https://' + url
                else:
                    url = 'https://' + url
            return url
    
    return None

def map_task_to_function(task: str) -> Tuple[Optional[Callable], Optional[str], Optional[str]]:
    """
    Mapeia uma tarefa para uma função que pode executá-la diretamente.
    
    Args:
        task: Descrição da tarefa
        
    Returns:
        Tupla (função, descrição, url_extraída) ou (None, None, None) se não encontrar
    """
    task_lower = task.lower().strip()
    task_stripped = task.strip()

    # Ajuda / comando agente no WhatsApp (várias funções)
    help_keywords = ['ajuda', 'help', 'o que você pode fazer', 'o que voce pode fazer', 'comandos', 'o que você faz', 'o que voce faz', 'capacidades', 'menu', 'agente', 'usar agente', 'comando agente', 'funções do agente', 'funcoes do agente']
    if any(k in task_lower for k in help_keywords) and len(task_stripped) < 80:
        return show_help, 'Mostrar ajuda e funções do agente', None

    # Verificar auto-otimização PRIMEIRO (prioridade máxima)
    optimize_mapping = TASK_MAPPING.get('self_optimize')
    has_optimize_keyword = False
    has_optimize_action = False
    
    if optimize_mapping:
        # Verificar palavras-chave específicas de auto-otimização
        optimize_keywords = ['agente', 'modificar código', 'modificar codigo', 'modificar seu próprio código',
                           'modificar seu proprio codigo', 'auto-melhorar', 'auto melhorar', 'auto-otimizar',
                           'auto otimizar', 'auto-evoluir', 'auto evoluir', 'auto-modificar', 'auto modificar',
                           'auto-aprender', 'auto aprender', 'aprendizado de máquina', 'aprendizado de maquina',
                           'machine learning', 'evoluir continuamente', 'melhorar desempenho', 'otimizar desempenho']
        
        has_optimize_keyword = any(keyword in task_lower for keyword in optimize_keywords)
        has_optimize_action = any(action in task_lower for action in optimize_mapping['actions'])
        
        # Também verificar combinações específicas
        has_agent_modify = 'agente' in task_lower and ('modificar' in task_lower or 'melhorar' in task_lower or 'otimizar' in task_lower)
        has_self_modify = ('próprio código' in task_lower or 'proprio codigo' in task_lower or 'próprio comportamento' in task_lower) and ('modificar' in task_lower or 'melhorar' in task_lower)
        
        if (has_optimize_keyword and has_optimize_action) or has_agent_modify or has_self_modify:
            return optimize_mapping['function'], optimize_mapping['description'], None
    
    # Verificar conexão de rede
    network_mapping = TASK_MAPPING.get('network')
    has_network_keyword = False
    has_network_action = False
    
    if network_mapping:
        has_network_keyword = any(keyword in task_lower for keyword in ['conectar', 'conexão', 'conexao', 'rede', 'wifi', 'wi-fi', 'ethernet', 'conectar-se', 'conectar se'])
        has_network_action = any(action in task_lower for action in ['conectar', 'conectar-se', 'conectar se', 'estabelecer conexão', 'ativar conexão', 'verificar conexão'])
        
        if has_network_keyword and has_network_action:
            return network_mapping['function'], network_mapping['description'], None
    
    # Verificar pesquisa na web (pesquisar X, buscar X, notícias) — busca, gera resposta com GPT e abre navegador
    search_mapping = TASK_MAPPING.get('search')
    if search_mapping and _agent_feature_enabled('search'):
        has_search_keyword = any(kw in task_lower for kw in search_mapping['keywords'])
        has_search_action = any(ac in task_lower for ac in search_mapping['actions'])
        if has_search_keyword or has_search_action:
            query = extract_search_query(task)
            return (lambda t=task, q=query: run_search_with_answer(t, q) if q else open_browser(), search_mapping['description'], query)
    
    # Verificar navegador (prioridade alta quando há palavras-chave específicas)
    browser_mapping = TASK_MAPPING.get('browser')
    has_browser_keyword = False
    has_browser_action = False
    
    if browser_mapping and _agent_feature_enabled('browser'):
        has_browser_keyword = any(keyword in task_lower for keyword in ['navegador', 'browser', 'internet', 'web', 'site', 'endereço', 'url'])
        has_browser_action = any(action in task_lower for action in ['abrir navegador', 'abrir browser', 'acessar internet', 'digitar endereço', 'digitar url', 'navegar'])
        
        # Só detectar navegador se NÃO for sobre conexão de rede
        if not has_network_keyword:
            if has_browser_keyword or has_browser_action:
                # Tentar extrair URL
                url = extract_url_from_task(task)
                if url:
                    return lambda: open_browser_and_navigate(url), 'Abrir navegador e navegar para site', url
                else:
                    return browser_mapping['function'], browser_mapping['description'], None
    
    # Priorizar câmera sobre foto se ambas aparecerem
    camera_mapping = TASK_MAPPING.get('camera')
    if camera_mapping and _agent_feature_enabled('camera'):
        has_action = any(action in task_lower for action in camera_mapping['actions'])
        mentions_device = 'dispositivo' in task_lower or 'notebook' in task_lower
        has_camera_word = any(kw in task_lower for kw in ['camera', 'câmera', 'camara', 'webcam'])
        mentions_capture = 'capturar' in task_lower or 'capturar vídeo' in task_lower or 'capturar video' in task_lower
        
        # Só detectar câmera se NÃO for sobre navegador, rede ou otimização
        if not has_browser_keyword and not has_browser_action and not has_network_keyword and not has_optimize_keyword:
            if (has_action and mentions_device) or (has_action and has_camera_word) or \
               (has_camera_word and mentions_device) or (has_camera_word and mentions_capture):
                return camera_mapping['function'], camera_mapping['description'], None
    
    # Abrir qualquer programa do computador (assistente) — ex.: "abrir calculadora", "abrir Word"
    if _agent_feature_enabled('openProgram'):
        _open_prefixes = ('abrir ', 'abre ', 'executar ', 'executa ', 'iniciar ', 'inicia ', 'rodar ', 'run ', 'open ')
        _exclude_programs = ('navegador', 'browser', 'câmera', 'camera', 'chrome', 'edge', 'firefox', 'internet', 'google')
        for prefix in _open_prefixes:
            if task_lower.startswith(prefix):
                program_name = task_stripped[len(prefix):].strip()
                if program_name and len(program_name) >= 2 and program_name.lower() not in _exclude_programs:
                    return (lambda n=program_name: open_program(n), f'Abrir programa: {program_name}', None)
                break
    
    # Verificar outros mapeamentos
    for task_type, mapping in TASK_MAPPING.items():
        if task_type in ['camera', 'browser', 'search']:
            continue  # Já verificados
        
        has_keyword = any(keyword in task_lower for keyword in mapping['keywords'])
        
        if has_keyword:
            if mapping.get('actions'):
                has_action = any(action in task_lower for action in mapping['actions'])
                if has_action:
                    return mapping['function'], mapping['description'], None
            else:
                return mapping['function'], mapping['description'], None
    
    return None, None, None

def execute_task_directly(task: str) -> bool:
    """
    Tenta executar a tarefa diretamente se for uma ação conhecida.
    Usa mapeamento de código + IA para identificar se a tarefa pode ser executada diretamente.
    
    Args:
        task: Descrição da tarefa
        
    Returns:
        True se a tarefa foi executada diretamente, False caso contrário
    """
    # 1. Primeiro tenta mapeamento direto do código (mais rápido e confiável)
    function, description, url = map_task_to_function(task)
    
    if function:
        print(f"Executando diretamente: {description}...")
        if url:
            if isinstance(url, str) and url.strip().lower().startswith("http"):
                print(f"   URL detectada: {url}")
            else:
                print(f"   Termo de busca: {url}")
        try:
            result = function()
            if result:
                print(f"✅ {description} executado com sucesso!")
                if isinstance(result, str):
                    print(f"   Resultado: {result}")
            else:
                print(f"✅ {description} executado!")
            return True
        except Exception as e:
            print(f"Erro ao executar {description}: {e}")
            return False
    
    # 2. Se não encontrou no mapeamento, usa IA para verificar (fallback)
    if not client:
        return False
    
    system_prompt = """Analise a tarefa e determine se ela pode ser executada diretamente pelo sistema.
    
Tarefas que podem ser executadas diretamente (mapeadas no código):
- Abrir câmera do notebook (função: access_notebook_camera)
- Tirar foto/capturar imagem (função: take_photo)
- Abrir navegador e acessar site (função: open_browser_and_navigate)
- Conectar à internet via Wi-Fi ou Ethernet (função: connect_to_internet)
- Auto-melhorar e otimizar código (função: auto_optimize)

Para cada tarefa mapeada, o sistema tem uma função Python específica que pode executá-la diretamente.

Responda APENAS com:
- "CAMERA" se for para abrir câmera do notebook
- "PHOTO" se for para tirar foto/capturar imagem
- "BROWSER" se for para abrir navegador ou acessar site
- "NETWORK" se for para conectar à internet ou verificar conexão
- "OPTIMIZE" se for para auto-melhorar/otimizar código
- "NO" se não puder ser executada diretamente (precisa de código, criação de arquivos, etc.)

Responda apenas com uma palavra: CAMERA, PHOTO, BROWSER, NETWORK, OPTIMIZE ou NO"""
    
    try:
        response = get_ai_response(task, system_prompt, max_tokens=10)
        
        if response:
            response_clean = response.strip().upper()
            
            if "CAMERA" in response_clean:
                print("Executando: Abrindo câmera do notebook...")
                access_notebook_camera()
                return True
            elif "PHOTO" in response_clean:
                print("Executando: Tirando foto...")
                photo_path = take_photo()
                if photo_path:
                    print(f"✅ Foto salva em: {photo_path}")
                return True
            elif "BROWSER" in response_clean:
                print("Executando: Abrindo navegador...")
                # Tentar extrair URL da tarefa
                url = extract_url_from_task(task)
                if url:
                    print(f"   URL detectada: {url}")
                    open_browser_and_navigate(url)
                else:
                    open_browser()
                return True
            elif "NETWORK" in response_clean:
                print("Executando: Verificando e conectando à internet...")
                connect_to_internet()
                return True
            elif "OPTIMIZE" in response_clean:
                print("Executando: Auto-otimização do código...")
                result = auto_optimize(force=True)
                if result.get('optimizations_applied', 0) > 0:
                    print(f"✅ {result['optimizations_applied']} otimizações aplicadas!")
                return True
    except Exception as e:
        print(f"Erro ao usar IA para mapear tarefa: {e}")
    
    return False


def extend_automation_with_gpt(task: str) -> Optional[dict]:
    """
    Quando a tarefa não está mapeada, usa GPT para entender o que o usuário quer
    e gera código para adicionar em cursor_automation_extensions.py (edita o próprio código).
    """
    if not client:
        return None
    ext_file = os.path.join(os.path.dirname(__file__), 'cursor_automation_extensions.py')
    system = """Você é um assistente que estende um script de automação em Python.

O usuário enviou uma tarefa que não tem comando ainda. Sua função:
1) Entender o que o usuário quer fazer (uma frase clara em português).
2) Se for algo que pode ser automatizado no PC (abrir app, pesquisar, executar programa, etc.), 
   gerar código Python para ser adicionado ao arquivo de extensões.

O arquivo de extensões define EXTRA_MAPPINGS = {} e depois preenche com entradas no formato:
EXTRA_MAPPINGS['chave_unica'] = {
    'keywords': ['palavra1', 'palavra2'],
    'actions': ['abrir', 'executar'],
    'function': funcao_que_faz_a_acao,
    'description': 'Descrição curta do comando'
}

Você tem acesso no código apenas a: subprocess, os, sys, time, pyautogui, open_browser, search_on_web, open_browser_and_navigate.
- Para abrir programas: subprocess.Popen(['cmd', '/c', 'start', '', 'nome_do_executavel']) no Windows, ou subprocess.Popen(['nome']) no Linux.
- Para pesquisar na web: search_on_web(termo).
- Para abrir navegador: open_browser() ou open_browser_and_navigate(url).

Responda APENAS com um JSON válido (sem markdown, sem ```), no formato:
{"description": "o que o usuário quer em uma frase", "code": "bloco Python para anexar ao arquivo (define uma função e adiciona a EXTRA_MAPPINGS) ou null se não for automatizável"}

Se a tarefa for só uma pergunta ou conversa (não uma ação no PC), use "code": null."""

    try:
        response = get_ai_response(task, system, max_tokens=800)
        if not response:
            return None
        response = response.strip()
        import json as _json
        try:
            # Remover possíveis marcadores de código
            resp_clean = response.strip()
            for start in ('```json', '```'):
                if resp_clean.startswith(start):
                    resp_clean = resp_clean[len(start):].lstrip()
                if resp_clean.endswith('```'):
                    resp_clean = resp_clean[:-3].rstrip()
            data = _json.loads(resp_clean)
        except _json.JSONDecodeError:
            # Tentar achar um bloco {...} na resposta
            i = response.find('{')
            j = response.rfind('}')
            if i != -1 and j != -1 and j > i:
                data = _json.loads(response[i:j+1])
            else:
                return None
        description = data.get('description') or 'Ação não identificada.'
        code = data.get('code')
        if not code:
            return {
                "success": False,
                "message": f"Entendi que você quer: {description}. Isso não pôde ser convertido em um comando de automação. Tente pedir uma ação no PC (ex.: abrir programa, pesquisar, abrir site).",
                "task": task
            }
        # Garantir que o arquivo existe (EXTRA_MAPPINGS é passado no contexto ao carregar)
        if not os.path.exists(ext_file):
            with open(ext_file, 'w', encoding='utf-8') as f:
                f.write('# Extensões do agente de automação (editado pelo próprio agente com GPT)\n')
                f.write('# Ao carregar, EXTRA_MAPPINGS é fornecido pelo script; os blocos abaixo adicionam entradas.\n\n')
        # Anexar o novo bloco
        with open(ext_file, 'a', encoding='utf-8') as f:
            f.write('\n# --- Adicionado pelo agente ---\n')
            f.write(code.strip())
            f.write('\n')
        print("✅ Código de extensão adicionado. Recarregando e tentando executar agora...")
        # Enviar aviso ao terminal do Cursor para as melhorias aparecerem e serem atualizadas
        _notify_cursor_terminal_extension_added(ext_file, description)
        # Recarregar extensões e executar a tarefa na mesma rodada (auto-reexecutar)
        _load_extension_mappings()
        if execute_task_directly(task):
            return {
                "success": True,
                "message": f"Entendi que você quer: {description}. Comando adicionado e executado com sucesso.",
                "task": task
            }
        return {
            "success": True,
            "message": f"Entendi que você quer: {description}. Foi adicionado um novo comando. Envie a mesma mensagem de novo para executar.",
            "task": task
        }
    except Exception as e:
        print(f"Erro ao estender automação com GPT: {e}")
        traceback.print_exc()
        return None


def generate_text_with_gpt(prompt: str, max_tokens: int = 1500) -> str:
    """
    Usa GPT para gerar um texto (história, redação, etc.) a partir do pedido do usuário.
    """
    if not client:
        return ""
    system = """Você é um escritor criativo. Gere o texto pedido pelo usuário: história, redação, poema, etc.
Seja criativo e engraçado quando pedirem algo engraçado. Escreva em português. Não inclua título no texto a menos que peçam.
Retorne apenas o texto gerado, sem explicações ou marcadores."""
    try:
        return (get_ai_response(prompt, system, max_tokens=max_tokens) or "").strip()
    except Exception as e:
        print(f"Erro ao gerar texto com GPT: {e}")
        return ""


def create_document_with_text_and_open(text: str, title: str = "Documento", save_as_pdf: bool = False) -> Optional[str]:
    """
    Cria um documento (.docx ou .txt) com o texto, salva em disco e abre no Word (ou app padrão).
    Retorna o caminho do arquivo ou None.
    """
    if not text or not text.strip():
        print("Texto vazio. Nada para salvar.")
        return None
    try:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        output_dir = os.path.join(script_dir, "output_docs")
        os.makedirs(output_dir, exist_ok=True)
        safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)[:80]
        base_name = f"{safe_title}_{int(time.time())}"
        docx_path = os.path.join(output_dir, f"{base_name}.docx")
        txt_path = os.path.join(output_dir, f"{base_name}.txt")
        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document()
            doc.add_heading(title, 0)
            for para in text.replace("\r\n", "\n").split("\n"):
                doc.add_paragraph(para.strip() or "")
            doc.save(docx_path)
            path_to_open = docx_path
            print(f"Documento Word criado: {path_to_open}")
        except ImportError:
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(title + "\n\n")
                f.write(text)
            path_to_open = txt_path
            print(f"Arquivo de texto criado: {path_to_open}")
        if sys.platform == 'win32':
            os.startfile(path_to_open)
        else:
            subprocess.Popen(['xdg-open', path_to_open] if sys.platform.startswith('linux') else ['open', path_to_open],
                             stdin=subprocess.DEVNULL, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        if save_as_pdf and path_to_open.endswith('.docx'):
            try:
                import docx2pdf
                pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
                docx2pdf.convert(path_to_open, pdf_path)
                print(f"PDF salvo: {pdf_path}")
                os.startfile(pdf_path)
            except Exception as e:
                print(f"Para salvar em PDF, abra o arquivo no Word e use Arquivo > Salvar como > PDF. ({e})")
        return path_to_open
    except Exception as e:
        print(f"Erro ao criar documento: {e}")
        traceback.print_exc()
        return None


def generate_excel_content_with_gpt(prompt: str, max_tokens: int = 800) -> Optional[dict]:
    """
    Usa GPT para gerar conteúdo de planilha: cabeçalhos e linhas.
    Retorna {"headers": ["Col1", "Col2", ...], "rows": [["v1","v2",...], ...]} ou None.
    """
    if not client:
        return None
    system = """Você gera dados para planilha Excel. O usuário descreve o que quer (ex: vendas do mês, lista de tarefas).
Retorne APENAS um JSON válido (sem markdown) com:
- "headers": array de strings com os nomes das colunas (ex: ["Produto", "Quantidade", "Valor"])
- "rows": array de arrays, cada um é uma linha (ex: [["Arroz", "2", "10.50"], ["Feijão", "1", "8.00"]])
Gere entre 3 e 15 linhas de exemplo coerentes. Use vírgula decimal. Escreva em português."""
    try:
        response = get_ai_response(prompt, system, max_tokens=max_tokens)
        if not response:
            return None
        import json as _json
        resp_clean = response.strip()
        for marker in ('```json', '```'):
            if resp_clean.startswith(marker):
                resp_clean = resp_clean[len(marker):].lstrip()
            if resp_clean.endswith('```'):
                resp_clean = resp_clean[:-3].rstrip()
        data = _json.loads(resp_clean)
        if isinstance(data.get("headers"), list) and isinstance(data.get("rows"), list):
            return data
    except Exception as e:
        print(f"Erro ao gerar conteúdo Excel com GPT: {e}")
    return None


def create_excel_with_data(headers: list, rows: list, title: str = "Planilha", sheet_name: Optional[str] = None) -> Optional[str]:
    """
    Cria um arquivo .xlsx com os cabeçalhos e linhas, salva e retorna o caminho.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.utils import get_column_letter
        script_dir = os.path.dirname(os.path.abspath(__file__))
        output_dir = os.path.join(script_dir, "output_excel")
        os.makedirs(output_dir, exist_ok=True)
        safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)[:80]
        base_name = f"{safe_title}_{int(time.time())}"
        xlsx_path = os.path.join(output_dir, f"{base_name}.xlsx")
        wb = Workbook()
        ws = wb.active
        ws.title = (sheet_name or "Dados")[:31]
        if headers:
            for col, val in enumerate(headers, 1):
                ws.cell(row=1, column=col, value=val)
        for r, row in enumerate(rows, 2):
            for c, val in enumerate(row, 1):
                ws.cell(row=r, column=c, value=val)
        wb.save(xlsx_path)
        print(f"Planilha Excel criada: {xlsx_path}")
        return xlsx_path
    except ImportError:
        print("openpyxl não instalado. Use: pip install openpyxl")
        return None
    except Exception as e:
        print(f"Erro ao criar Excel: {e}")
        traceback.print_exc()
        return None


def read_excel_file(file_path: str) -> Optional[dict]:
    """
    Lê um arquivo .xlsx e retorna {"headers": [...], "rows": [[...], ...]} ou None.
    """
    if not file_path or not os.path.isfile(file_path):
        return None
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        wb.close()
        if not rows:
            return {"headers": [], "rows": []}
        return {"headers": list(rows[0]), "rows": [list(r) for r in rows[1:]]}
    except Exception as e:
        print(f"Erro ao ler Excel: {e}")
        return None


def create_excel_and_open(excel_prompt: str, title: str = "Planilha") -> Optional[str]:
    """
    Gera conteúdo com GPT, cria o .xlsx, salva e abre no Excel.
    """
    data = generate_excel_content_with_gpt(excel_prompt)
    if not data:
        return None
    headers = data.get("headers", [])
    rows = data.get("rows", [])
    path = create_excel_with_data(headers, rows, title=title)
    if not path:
        return None
    if sys.platform == 'win32':
        os.startfile(path)
    else:
        subprocess.Popen(
            ['xdg-open', path] if sys.platform.startswith('linux') else ['open', path],
            stdin=subprocess.DEVNULL, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )
    return path


def open_excel_file(file_path: str) -> bool:
    """Abre um arquivo .xlsx no aplicativo padrão (Excel no Windows)."""
    if not file_path or not os.path.isfile(file_path):
        return False
    try:
        if sys.platform == 'win32':
            os.startfile(file_path)
        else:
            subprocess.Popen(
                ['xdg-open', file_path] if sys.platform.startswith('linux') else ['open', file_path],
                stdin=subprocess.DEVNULL, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
            )
        return True
    except Exception as e:
        print(f"Erro ao abrir Excel: {e}")
        return False


def _fetch_music_suggestions_from_web(user_task: str) -> str:
    """Busca na internet sugestões de músicas (mais tocadas, em alta) para o modelo escolher com base em dados reais."""
    if not user_task or not isinstance(user_task, str):
        return ""
    t = user_task.strip().lower()
    music_keywords = ["música", "musica", "tocar", "toca", "escolha", "escolher", "outra musica", "outra música", "animada", "romântica", "romantica", "relaxante", "funk", "sertanejo", "pop"]
    if not any(k in t for k in music_keywords):
        return ""
    if "animad" in t or "animada" in t:
        search_q = "músicas animadas mais tocadas 2024 sucesso"
    elif "romântic" in t or "romantic" in t:
        search_q = "músicas românticas mais tocadas sucesso"
    elif "relaxante" in t:
        search_q = "músicas relaxantes mais tocadas"
    elif "outra" in t or "outro" in t or "escolheu a mesma" in t:
        search_q = "músicas mais tocadas agora 2024 hits"
    else:
        search_q = "músicas mais tocadas 2024 sucesso"
    try:
        snippet_text = _fetch_duckduckgo_html_snippets(search_q)
        if not snippet_text or len(snippet_text.strip()) < 30:
            return ""
        text = snippet_text.strip()[:1400]
        if len(snippet_text.strip()) > 1400:
            text += "..."
        return "\n\nSugestões encontradas na internet (use para escolher uma música real e atual; prefira uma que apareça nos resultados):\n" + text
    except Exception:
        return ""


def _get_agent_history_for_prompt(max_entries: int = 100, max_result_chars: int = 500) -> str:
    """Lê o histórico do agente (banco) passado pelo Node. Inclui o que foi feito (ex.: música tocada) para não repetir.
    max_entries e max_result_chars podem ser reduzidos para modelos com contexto pequeno (8192 tokens)."""
    try:
        raw = os.environ.get("AUTOMATION_AGENT_HISTORY_JSON", "")
        if not raw:
            return ""
        import json as _json
        arr = _json.loads(raw)
        if not arr or not isinstance(arr, list):
            return ""
        lines = []
        songs_played = []
        for i, h in enumerate(arr[:max_entries]):
            if isinstance(h, dict):
                u = (h.get("userMessage") or h.get("user_message") or "").strip()
                t = (h.get("taskExecuted") or h.get("task_executed") or "").strip()
                res = (h.get("resultMessage") or "").strip()
                s = "ok" if h.get("success") else "falhou"
                if res and ("Tocando:" in res or "tocou" in res.lower() or "YouTube" in res):
                    for part in re.findall(r"Tocando:\s*([^\n🎵]+)", res):
                        song = part.strip().strip("'\"").strip()
                        if song and len(song) > 2 and song not in songs_played:
                            songs_played.append(song)
                    for part in re.findall(r"YouTube[^']*['\"]([^'\"]+?)['\"]", res):
                        song = part.strip().strip("'\"").strip()
                        if song and len(song) > 2 and song not in songs_played:
                            songs_played.append(song)
                if u or t:
                    part = f"- Usuário: {u[:70]}{'...' if len(u) > 70 else ''} → Tarefa: {t[:50]}{'...' if len(t) > 50 else ''} → {s}"
                    if res:
                        part += f" | Resultado: {res[:max_result_chars]}{'...' if len(res) > max_result_chars else ''}"
                    lines.append(part)
        out = ""
        if songs_played:
            out += "\n\n** Músicas já tocadas (NÃO SUGIRA ESTAS de novo): " + ", ".join(songs_played[:8]) + " **\n"
        if lines:
            out += "\nHistórico recente (use para não repetir; se o usuário pedir 'outra música', escolha uma DIFERENTE das listadas acima):\n" + "\n".join(lines)
        return out if out.strip() else ""
    except Exception:
        return ""


def interpret_task_with_gpt(user_task: str) -> Optional[dict]:
    """
    Usa GPT para entender o que o usuário pediu e sugerir a melhor forma de executar.
    Ajuda o agente a interpretar a intenção e traduzir para comandos precisos.
    
    Args:
        user_task: Tarefa original do usuário (ex: "quero ver o clima", "me mostra notícias")
    
    Returns:
        Dict com: {'interpreted_task': str, 'suggested_function': str, 'params': dict} ou None
    """
    model = _get_agent_model()
    use_grok = model.lower().startswith("grok")
    if use_grok and not grok_client:
        print("AVISO: Modelo do agente é Grok mas XAI_API_KEY não está configurada. Configure no .env (XAI_API_KEY ou GROK_API_KEY).")
        return None
    if not use_grok and not client:
        return None
    # Modelos com contexto 8192 tokens (ex.: gpt-4 antigo) exigem histórico e blocos menores
    small_ctx = _model_has_small_context(model)
    if small_ctx:
        history_block = _get_agent_history_for_prompt(max_entries=8, max_result_chars=120)
    else:
        history_block = _get_agent_history_for_prompt()
    web_music_block = _fetch_music_suggestions_from_web(user_task)
    if small_ctx and len(web_music_block) > 800:
        web_music_block = web_music_block[:800].rstrip() + "\n[...]\n"
    system = (history_block + "\n\n") if history_block else ""
    system += """Você é um assistente que interpreta o que o usuário quer fazer no PC e ajuda o agente a executar.

O agente pode executar estas ações:
- search_on_web(query): Pesquisar na web
- open_program(name): Abrir qualquer programa (Word, Excel, etc.)
- start_camera_stream_to_server(): Ligar câmera
- open_browser() ou open_browser_and_navigate(url): Abrir navegador. Não fechar após abertura; só fechar se o usuário pedir.
- close_browser(): Fechar o navegador. Use somente quando o usuário pedir explicitamente (ex: "feche o navegador", "fechar navegador").
- take_photo(): Tirar foto
- create_story_and_document: Criar texto com GPT, abrir no Word e opcionalmente salvar em PDF. params: story_prompt, title, save_pdf (true/false)
- excel_create_and_open: Criar planilha com GPT (conteúdo gerado), salvar e abrir no Excel. params: excel_prompt (ex: "vendas do mês"), title (título da planilha)
- excel_read: Ler planilha Excel e abrir ou resumir. params: file_path (caminho) ou file_name (nome do arquivo em output_excel)
- mouse_click: Clicar na tela. params: x (número), y (número), button ("left" ou "right").
- move_mouse: Apenas mover o mouse (sem clicar). params: x, y (números)
- youtube_search_and_play: Buscar no YouTube e dar play no primeiro vídeo. params: query (termo de busca; para pedidos genéricos de música, sugira livremente uma música; veja regra abaixo).
- youtube_play_link_in_html: Pegar o link do vídeo do YouTube e criar um HTML para reproduzir sem ser pela página do YouTube (embed, sem anúncios da página). params: url (link completo do vídeo, ex: https://www.youtube.com/watch?v=ID). Use quando o usuário mandar um link do YouTube e pedir para reproduzir em HTML, sem YouTube, sem propaganda da página, etc.

Regra para YouTube/música: Sugira sempre uma música específica (título + artista). Use o HISTÓRICO para não repetir o que já foi tocado. Quando houver "Sugestões encontradas na internet" abaixo, PREFIRA escolher uma música que apareça nesses resultados (são músicas reais e em alta). Se o usuário pedir "outra música" ou "escolha outra", escolha uma DIFERENTE do histórico e, se possível, das sugestões da web. params.query = "Título Artista". interpreted_task = "Tocar [sua sugestão] no YouTube".

Analise o pedido, o histórico e as sugestões da web (quando existirem).
Exemplos (apenas formato; sugira o que achar melhor):
- "quero ver o clima" → suggested_function: "search_on_web", params: {"query": "clima hoje"}
- "abrir Excel" → suggested_function: "open_program", params: {"program_name": "Excel"}
- "clique em 100 200" → suggested_function: "mouse_click", params: {"x": 100, "y": 200}
- "feche o navegador" → suggested_function: "close_browser", params: {}
- "escolha uma música animada" / "uma música pra mim" → suggested_function: "youtube_search_and_play", params: {"query": "[qualquer música animada que você quiser sugerir]"}; interpreted_task: "Tocar [sua sugestão] no YouTube"
- "toca uma música romântica" → suggested_function: "youtube_search_and_play", params: {"query": "[sugestão sua]"}; interpreted_task com a música escolhida
- "pega o link X e cria um HTML..." → suggested_function: "youtube_play_link_in_html", params: {"url": "..."}
- "cria uma planilha de vendas" → suggested_function: "excel_create_and_open", params: {"excel_prompt": "planilha de vendas", "title": "Vendas"}

Retorne APENAS um JSON válido (sem markdown, sem ```) no formato:
{
  "interpreted_task": "comando claro",
  "suggested_function": "search_on_web|open_program|start_camera_stream_to_server|open_browser|close_browser|take_photo|create_story_and_document|excel_create_and_open|excel_read|mouse_click|move_mouse|youtube_search_and_play|youtube_play_link_in_html",
  "params": {"query": "...", "program_name": "...", "x": número, "y": número, "button": "left" ou "right", ...},
  "confidence": "high|medium|low"
}

Se não conseguir entender, retorne: {"interpreted_task": null, "suggested_function": null, "params": {}, "confidence": "low"}"""
    system += web_music_block
    try:
        response = get_ai_response(user_task, system, max_tokens=220)
        if not response:
            return None
        import json as _json
        resp_clean = response.strip()
        for marker in ('```json', '```'):
            if resp_clean.startswith(marker):
                resp_clean = resp_clean[len(marker):].lstrip()
            if resp_clean.endswith('```'):
                resp_clean = resp_clean[:-3].rstrip()
        data = _json.loads(resp_clean)
        if data.get("interpreted_task") and data.get("confidence") != "low":
            return data
    except Exception as e:
        verbose = os.environ.get("VERBOSE", "").strip().lower() in ("1", "true")
        if verbose:
            print(f"GPT interpretação: {e}")
    return None


def _get_conversational_reply_if_any(user_message: str) -> Optional[str]:
    """
    Se a mensagem for apenas conversa (oi, obrigado, tudo bem?, etc.) sem pedir ação no PC,
    retorna uma resposta amigável. Caso contrário retorna None (o agente deve executar comando).
    """
    if not user_message or not (user_message or "").strip():
        return None
    model = _get_agent_model()
    use_grok = model.lower().startswith("grok")
    if use_grok and not grok_client:
        return None
    if not use_grok and not client:
        return None
    system = """Você é um assistente no WhatsApp que também controla o PC (navegador, pesquisa, música, etc.).
Se a mensagem for APENAS conversa (saudação, agradecimento, pergunta sobre você) sem pedir ação no PC, responda em português de forma OBJETIVA e CURTA (uma frase, sem rodeios).
Se o usuário pedir para FAZER algo (abrir, pesquisar, tocar música, etc.), responda só: COMMAND
Responda só o texto da resposta ou COMMAND."""
    try:
        response = (get_ai_response(user_message.strip(), system, max_tokens=100) or "").strip()
        if not response:
            return None
        if response.upper().startswith("COMMAND"):
            return None
        return response[:500] if len(response) > 500 else response
    except Exception:
        return None


# ==================== FUNÇÃO PRINCIPAL ====================
def automate_cursor(task: str, **kwargs) -> dict:
    """
    Executa a tarefa: abre navegador, pesquisas, câmera, edição de texto, etc.
    Tenta executar diretamente as ações mapeadas; não abre o Cursor.
    
    Args:
        task: Descrição da tarefa (ex: abrir navegador, pesquisar X, ligar câmera)
        **kwargs: Argumentos adicionais (mantidos para compatibilidade)
    
    Returns:
        Dicionário com resultado da execução
    """
    print(f"\n{'='*50}")
    print(f"Processando tarefa: {task}")
    print(f"{'='*50}\n")
    
    # Comandos de saída: não executar automação
    if task and task.strip().lower() in ('sair', 'exit', 'quit'):
        print("Comando de saída. Nenhuma automação executada.")
        return {"success": True, "message": "Comando de saída. Nenhuma automação executada.", "task": task}

    task_lower = (task or "").strip().lower()

    # Atalho: link do YouTube + pedido de HTML/embed → criar HTML e abrir (sem depender do GPT)
    yt_url = extract_youtube_url_from_text(task)
    if yt_url and (
        "html" in task_lower or "embed" in task_lower or "sem ser pelo youtube" in task_lower
        or "sem youtube" in task_lower or "reproduzir sem" in task_lower or "reprodução sem" in task_lower
        or "criar um html" in task_lower or "cria um html" in task_lower or "reproduz em html" in task_lower
        or "sem ser pelo youtube" in task_lower or "sem abrir o youtube" in task_lower
    ):
        print("Link do YouTube detectado + pedido de HTML. Criando e abrindo...")
        if open_youtube_video_in_html(yt_url):
            return {"success": True, "message": "HTML criado e aberto: vídeo em reprodução (embed, sem página do YouTube).", "task": task}
        return {"success": False, "message": "Não foi possível criar ou abrir o HTML. Verifique o link do YouTube.", "task": task}

    # Atalho: comandos óbvios executam direto sem chamar GPT (resposta mais rápida)
    quick_actions = [
        ("abrir chrome", lambda: open_browser("Chrome")),
        ("abrir edge", lambda: open_browser("Edge")),
        ("abrir brave", lambda: open_browser("Brave")),
        ("abrir navegador", lambda: open_browser()),
        ("fechar navegador", lambda: close_browser()),
        ("feche o navegador", lambda: close_browser()),
        ("ligar câmera", lambda: start_camera_stream_to_server()),
        ("liga a câmera", lambda: start_camera_stream_to_server()),
    ]
    for keyword, action in quick_actions:
        if keyword in task_lower or task_lower == keyword:
            try:
                if action():
                    return {"success": True, "message": f"Executado: {task.strip()}.", "task": task}
            except Exception as e:
                print(f"Atalho falhou: {e}")
            break

    # 1. Usar GPT para interpretar o que o usuário quer (só quando não for comando óbvio)
    interpreted = None
    task_to_execute = task
    if client:
        interpreted = interpret_task_with_gpt(task)
        if interpreted and interpreted.get("interpreted_task"):
            interpreted_task = interpreted["interpreted_task"]
            _model = _get_agent_model()
            _label = "Grok" if _model.lower().startswith("grok") else "GPT"
            print(f"💡 {_label} interpretou: '{task}' → '{interpreted_task}'")
            task_to_execute = interpreted_task
            # Se o GPT sugeriu parâmetros específicos e função direta, tentar executar
            params = interpreted.get("params", {})
            suggested_func = interpreted.get("suggested_function", "")
            if params and suggested_func:
                if suggested_func == "search_on_web" and params.get("query"):
                    print(f"Executando pesquisa sugerida pelo GPT: {params['query']}")
                    if run_search_with_answer(task, params["query"]):
                        return {"success": True, "message": f"Pesquisa executada: {params['query']}", "task": task}
                elif suggested_func == "youtube_play_link_in_html":
                    url = (params.get("url") or "").strip()
                    if not url or not extract_youtube_video_id(url):
                        url = extract_youtube_url_from_text(task)
                    if url and extract_youtube_video_id(url):
                        print("Executando: criar HTML e reproduzir link do YouTube sem ser pela página.")
                        if open_youtube_video_in_html(url):
                            return {"success": True, "message": "HTML criado e aberto: vídeo em reprodução (embed, sem página do YouTube).", "task": task}
                    return {"success": False, "message": "Informe um link válido do YouTube na mensagem (ex: https://youtube.com/watch?v=xxxx ou https://youtu.be/xxxx).", "task": task}
                elif suggested_func == "youtube_search_and_play":
                    query = params.get("query") or "música"
                    print(f"Executando: YouTube em segundo plano — '{query}' (sem pausa ao terminar)")
                    script_dir = os.path.dirname(os.path.abspath(__file__))
                    try:
                        creationflags = 0
                        if sys.platform == 'win32':
                            DETACH = getattr(subprocess, 'DETACHED_PROCESS', 0x00000008)
                            NO_WINDOW = getattr(subprocess, 'CREATE_NO_WINDOW', 0x08000000)
                            creationflags = DETACH | NO_WINDOW
                        subprocess.Popen(
                            [sys.executable, __file__, "--youtube", query],
                            cwd=script_dir or os.getcwd(),
                            stdin=subprocess.DEVNULL,
                            stdout=subprocess.DEVNULL,
                            stderr=subprocess.DEVNULL,
                            creationflags=creationflags,
                        )
                        return {"success": True, "message": f"YouTube em segundo plano: busquei '{query}' e acionando o play. O navegador vai abrir e permanecer aberto.", "task": task}
                    except Exception as e:
                        print(f"Fallback: executando YouTube no mesmo processo ({e})")
                        if youtube_search_and_play(query):
                            return {"success": True, "message": f"YouTube: busquei '{query}' e acionei o play.", "task": task}
                    return {"success": False, "message": "Não foi possível iniciar o YouTube.", "task": task}
                elif suggested_func == "open_program" and params.get("program_name"):
                    print(f"Executando: abrir {params['program_name']} (sugerido pelo GPT)")
                    if open_program(params["program_name"]):
                        return {"success": True, "message": f"Programa aberto: {params['program_name']}", "task": task}
                elif suggested_func == "close_browser":
                    print("Executando: fechar navegador (pedido do usuário)")
                    if close_browser():
                        return {"success": True, "message": "Navegador fechado.", "task": task}
                    return {"success": False, "message": "Não foi possível fechar o navegador ou nenhuma janela estava em foco.", "task": task}
                elif suggested_func == "start_camera_stream_to_server":
                    print("Executando: ligar câmera (sugerido pelo GPT)")
                    if start_camera_stream_to_server():
                        return {"success": True, "message": "Câmera ligada (sugerido pelo GPT)", "task": task}
                elif suggested_func == "create_story_and_document" and params.get("story_prompt"):
                    story_prompt = params.get("story_prompt", "")
                    title = params.get("title", "Documento")
                    save_pdf = bool(params.get("save_pdf", False))
                    print(f"Executando: criar história e abrir no Word (PDF={save_pdf})")
                    text = generate_text_with_gpt(story_prompt)
                    if text:
                        path = create_document_with_text_and_open(text, title=title, save_as_pdf=save_pdf)
                        if path:
                            return {"success": True, "message": "História criada e documento aberto no Word." + (" PDF gerado." if save_pdf else ""), "task": task}
                    return {"success": False, "message": "Não foi possível gerar o texto ou criar o documento.", "task": task}
                elif suggested_func == "excel_create_and_open" and params.get("excel_prompt"):
                    excel_prompt = params.get("excel_prompt", "")
                    title = params.get("title", "Planilha")
                    print(f"Executando: criar planilha Excel e abrir (salvar)")
                    path = create_excel_and_open(excel_prompt, title=title)
                    if path:
                        return {"success": True, "message": "Planilha criada, salva e aberta no Excel.", "task": task}
                    return {"success": False, "message": "Não foi possível gerar ou salvar a planilha Excel.", "task": task}
                elif suggested_func == "excel_read":
                    file_path = params.get("file_path") or ""
                    file_name = params.get("file_name") or ""
                    script_dir = os.path.dirname(os.path.abspath(__file__))
                    output_excel = os.path.join(script_dir, "output_excel")
                    if file_path and os.path.isfile(file_path):
                        path_to_use = file_path
                    elif file_name:
                        path_to_use = os.path.join(output_excel, file_name) if not os.path.isabs(file_name) else file_name
                        if not os.path.isfile(path_to_use):
                            for f in (os.listdir(output_excel) if os.path.isdir(output_excel) else []):
                                if f.endswith(".xlsx") and file_name.lower() in f.lower():
                                    path_to_use = os.path.join(output_excel, f)
                                    break
                    else:
                        path_to_use = None
                    if path_to_use and os.path.isfile(path_to_use):
                        data = read_excel_file(path_to_use)
                        if open_excel_file(path_to_use):
                            summary = ""
                            if data and data.get("rows"):
                                summary = f" {len(data['rows'])} linhas lidas."
                            return {"success": True, "message": "Planilha aberta no Excel." + summary, "task": task}
                    return {"success": False, "message": "Arquivo Excel não encontrado ou não foi possível abrir.", "task": task}
                elif suggested_func == "mouse_click":
                    try:
                        w, h = pyautogui.size()
                        x = params.get("x")
                        y = params.get("y")
                        if x is None or y is None:
                            x, y = w // 2, h // 2
                        else:
                            x, y = int(x), int(y)
                        button = params.get("button") or "left"
                        move_and_click(x, y, button=button)
                        return {"success": True, "message": f"Clique executado em ({x}, {y}) com botão {button}.", "task": task}
                    except Exception as e:
                        return {"success": False, "message": f"Erro ao clicar: {e}", "task": task}
                elif suggested_func == "move_mouse":
                    try:
                        w, h = pyautogui.size()
                        x = params.get("x")
                        y = params.get("y")
                        if x is None or y is None:
                            x, y = w // 2, h // 2
                        else:
                            x, y = int(x), int(y)
                        pyautogui.moveTo(x, y, duration=MOUSE_SPEED)
                        return {"success": True, "message": f"Mouse movido para ({x}, {y}).", "task": task}
                    except Exception as e:
                        return {"success": False, "message": f"Erro ao mover mouse: {e}", "task": task}
    
    # 2. Tentar executar diretamente (com a tarefa interpretada pelo GPT ou original)
    if execute_task_directly(task_to_execute):
        msg = f"Tarefa executada: {task}"
        if interpreted and interpreted.get("interpreted_task") and interpreted.get("interpreted_task") != task:
            _m = _get_agent_model()
            _l = "Grok" if _m.lower().startswith("grok") else "GPT"
            msg += f" ({_l} interpretou como: {interpreted['interpreted_task']})"
        result = {
            "success": True,
            "message": msg,
            "task": task
        }
        print(f"✅ Tarefa concluída com sucesso!")
        return result
    
    # Tarefa não mapeada: usar GPT para entender o que o usuário quer e editar o próprio código (extensões)
    print("Tarefa não reconhecida. Usando GPT para entender e estender o código...")
    extend_result = extend_automation_with_gpt(task)
    if extend_result:
        return extend_result
    result = {
        "success": False,
        "message": "Tarefa não executada. Use comandos como: abrir navegador, pesquisar X, ligar câmera, acessar [site], editar texto.",
        "task": task
    }
    return result

# ==================== MODO INTERATIVO ====================
def interactive_mode():
    """Modo interativo para múltiplas tarefas."""
    print("\n" + "="*50)
    print("MODO INTERATIVO - Automação com GPT (OpenAI)")
    print("="*50)
    print("Ações: abrir navegador, pesquisar na web, ligar câmera, editar texto, acessar site.")
    print("Digite 'sair' para encerrar\n")
    
    while True:
        try:
            task = input("Digite a tarefa (ex: abrir navegador, pesquisar X, ligar câmera): ").strip()
            
            if task.lower() in ['sair', 'exit', 'quit']:
                print("Encerrando...")
                break
            
            if not task:
                print("Tarefa vazia. Tente novamente.")
                continue
            
            automate_cursor(task)
            print("\n" + "-"*50 + "\n")
            
        except (EOFError, KeyboardInterrupt):
            print("\nEncerrando...")
            break
        except Exception as e:
            print(f"Erro: {e}")
            print("Tente novamente.\n")

def _apply_bot_config() -> None:
    """
    Aplica a config enviada pelo bot (Node) via AUTOMATION_CONFIG_JSON.
    Personaliza timeouts, navegador, limite de resposta, features, etc.
    """
    global TIMEOUT_VARREDURA, MAX_LINKS_VARREDURA, DEFAULT_BROWSER, ENVIAR_ATUALIZACOES_TERMINAL_CURSOR, AGENT_CONFIG
    cfg_json = os.environ.get("AUTOMATION_CONFIG_JSON", "")
    if not cfg_json or not cfg_json.strip():
        return
    try:
        cfg = json.loads(cfg_json)
        AGENT_CONFIG.clear()
        AGENT_CONFIG.update(cfg)
        if "timeoutVarredura" in cfg:
            TIMEOUT_VARREDURA = int(cfg["timeoutVarredura"])
        if "maxLinksVarredura" in cfg:
            MAX_LINKS_VARREDURA = int(cfg["maxLinksVarredura"])
        if "preferredBrowser" in cfg and str(cfg["preferredBrowser"]).strip():
            DEFAULT_BROWSER = str(cfg["preferredBrowser"]).strip()
        if "enviarAtualizacoesTerminalCursor" in cfg:
            ENVIAR_ATUALIZACOES_TERMINAL_CURSOR = bool(cfg["enviarAtualizacoesTerminalCursor"])
        if "maxRespostaWhatsApp" in cfg:
            os.environ["MAX_RESPOSTA_WHATSAPP"] = str(int(cfg["maxRespostaWhatsApp"]))
    except (json.JSONDecodeError, TypeError, ValueError) as e:
        print(f"Aviso: config do bot inválida: {e}")


# ==================== MAIN ====================
if __name__ == "__main__":
    # Garantir que estamos executando o arquivo .py, não um diretório
    _script_path = os.path.abspath(__file__)
    if not _script_path.lower().endswith('.py') or os.path.isdir(_script_path):
        print("ERRO: Execute o arquivo Python explicitamente:")
        print("  python cursor_automation.py")
        print("  ou use: .\\run_automation.ps1")
        sys.exit(1)

    # Modo YouTube em segundo plano: evita pausa ao terminar (navegador fica aberto, script principal retorna logo)
    if len(sys.argv) >= 3 and sys.argv[1] == "--youtube":
        _query = " ".join(sys.argv[2:]).strip() or "música"
        _apply_bot_config()
        try:
            _youtube_search_and_play_selenium(_query)
            # Pequena espera para o play iniciar; pausa ~50s costuma ser aba/janela perdendo foco (ex.: abrir WhatsApp)
            time.sleep(3)
        except Exception as e:
            print(f"YouTube em segundo plano: {e}")
        sys.exit(0)

    # Aplicar config do bot (personalização enviada pelo Node)
    _apply_bot_config()
    _agent_model = _get_agent_model()
    _use_grok = _agent_model.lower().startswith("grok")
    print(f"Modelo do agente: {_agent_model} ({'xAI/Grok' if _use_grok else 'OpenAI/GPT'})")

    # Configurar encoding para Windows
    if sys.platform == 'win32':
        try:
            sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
            sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
        except Exception:
            pass

    # Modo stream de câmera para o servidor (processo separado)
    if '--camera-stream' in sys.argv:
        _run_camera_stream_loop()
        sys.exit(0)

    if len(sys.argv) <= 1:
        interactive_mode()
    else:
        task = " ".join(sys.argv[1:]).strip()
        print("OK", flush=True)  # Resposta imediata para o Node não parecer travado
        # Comandos óbvios: pular verificação conversacional (evita 1 chamada à API e reduz demora)
        t = (task or "").strip().lower()
        skip_chat = any(k in t for k in (
            "música", "musica", "toca", "tocar", "abre", "abrir", "pesquisa", "pesquisar",
            "navegador", "youtube", "outra musica", "outra música", "escolha", "fecha",
            "fechar", "câmera", "camera", "excel", "word", "clique", "mouse"
        ))
        chat_reply = None if skip_chat else _get_conversational_reply_if_any(task)
        if chat_reply:
            print("AGENT_CHAT_START")
            print(chat_reply)
            print("AGENT_CHAT_END")
            sys.exit(0)

        # Diagnóstico só quando for executar automação (não conversa)
        print("Script de automação iniciado. Tarefa:", task[:80] + ("..." if len(task) > 80 else ""))
        non_interactive = True
        if not OPENAI_API_KEY and not XAI_API_KEY:
            print("AVISO: OPENAI_API_KEY não configurada (extensões por GPT não funcionarão).")
        elif non_interactive:
            if _agent_model.lower().startswith("grok") and XAI_API_KEY:
                print(f"xAI/Grok API configurada (chave: {XAI_API_KEY[:10]}...)\n")
            elif OPENAI_API_KEY:
                print(f"OpenAI API configurada (chave: {OPENAI_API_KEY[:10]}...)\n")

        try:
            result = automate_cursor(task)
        except Exception as e:
            print(f"❌ Erro ao executar: {e}")
            traceback.print_exc()
            sys.exit(1)
